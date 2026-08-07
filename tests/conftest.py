"""Shared test helpers."""

from __future__ import annotations

import posixpath
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


@pytest.fixture
def convert(tmp_path):
    """Converts Markdown text and returns (path, result)."""
    from md2word.config import Config
    from md2word.converter import convert_text

    counter = {"n": 0}

    def _convert(markdown: str, **options):
        counter["n"] += 1
        config = Config(base_dir=str(tmp_path), **options)
        target = tmp_path / f"out{counter['n']}.docx"
        result = convert_text(markdown, str(target), config)
        return target, result

    return _convert


@pytest.fixture
def doc(convert):
    """Converts and returns the opened Document directly."""

    def _doc(markdown: str, **options):
        path, _result = convert(markdown, **options)
        return Document(str(path))

    return _doc


def paragraph_styles(document) -> list[str]:
    return [p.style.style_id for p in document.paragraphs]


def texts(document) -> list[str]:
    return [p.text for p in document.paragraphs]


def find_paragraph(document, needle: str):
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise AssertionError(f"no paragraph contains {needle!r}: {texts(document)}")


def numbering_of(paragraph) -> tuple[int, int] | None:
    """(numId, ilvl) of a list paragraph."""
    pPr = paragraph._p.find(w("pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(w("numPr"))
    if numPr is None:
        return None
    num_id = numPr.find(w("numId"))
    ilvl = numPr.find(w("ilvl"))
    return (
        int(num_id.get(w("val"))) if num_id is not None else -1,
        int(ilvl.get(w("val"))) if ilvl is not None else 0,
    )


def validate_package(path: Path) -> list[str]:
    """Structural check of the OPC package - returns the list of problems."""
    problems: list[str] = []
    z = zipfile.ZipFile(path)
    names = set(z.namelist())

    if z.testzip():
        problems.append("corrupt ZIP entry")

    for name in names:
        if name.endswith((".xml", ".rels")):
            try:
                etree.fromstring(z.read(name))
            except etree.XMLSyntaxError as exc:
                problems.append(f"malformed XML in {name}: {exc}")

    ct = etree.fromstring(z.read("[Content_Types].xml"))
    defaults = {e.get("Extension").lower() for e in ct.findall(f"{{{CT_NS}}}Default")}
    overrides = {e.get("PartName").lstrip("/") for e in ct.findall(f"{{{CT_NS}}}Override")}
    for name in names:
        if name == "[Content_Types].xml" or name.endswith(".rels"):
            continue
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if name not in overrides and ext not in defaults:
            problems.append(f"no content type for {name}")
    for part in overrides:
        if part not in names:
            problems.append(f"override without a part: {part}")

    for name in (n for n in names if n.endswith(".rels")):
        base = posixpath.dirname(posixpath.dirname(name))
        for rel in etree.fromstring(z.read(name)).findall(f"{{{REL_NS}}}Relationship"):
            if rel.get("TargetMode") == "External":
                continue
            resolved = posixpath.normpath(
                posixpath.join(base, rel.get("Target"))
            ).lstrip("/")
            if resolved not in names:
                problems.append(f"{name}: missing target -> {rel.get('Target')}")

    rels = etree.fromstring(z.read("word/_rels/document.xml.rels"))
    rids = {r.get("Id") for r in rels.findall(f"{{{REL_NS}}}Relationship")}
    document = etree.fromstring(z.read("word/document.xml"))
    used = set(document.xpath("//@r:id", namespaces={"r": R}))
    if used - rids:
        problems.append(f"unknown rIds: {sorted(used - rids)}")

    # Styles resolvable
    styles = etree.fromstring(z.read("word/styles.xml"))
    style_ids = {s.get(w("styleId")) for s in styles.findall(w("style"))}
    for part in ("word/document.xml", "word/footnotes.xml"):
        if part not in names:
            continue
        tree = etree.fromstring(z.read(part))
        for tag in ("pStyle", "rStyle", "tblStyle"):
            for node in tree.iter(w(tag)):
                value = node.get(w("val"))
                if value and value not in style_ids:
                    problems.append(f"{part}: unknown style {value!r}")

    # Numbering
    if "word/numbering.xml" in names:
        numbering = etree.fromstring(z.read("word/numbering.xml"))
        defined = {n.get(w("numId")) for n in numbering.findall(w("num"))}
        abstract = {n.get(w("abstractNumId")) for n in numbering.findall(w("abstractNum"))}
        for node in numbering.findall(w("num")):
            ref = node.find(w("abstractNumId"))
            if ref is not None and ref.get(w("val")) not in abstract:
                problems.append("num points at a missing abstractNum")
        children = [c.tag for c in numbering if isinstance(c.tag, str)]
        last_abstract = max(
            (i for i, t in enumerate(children) if t.endswith("}abstractNum")), default=-1
        )
        first_num = min(
            (i for i, t in enumerate(children) if t.endswith("}num")), default=len(children)
        )
        if last_abstract > first_num:
            problems.append("numbering.xml: abstractNum after num (schema violation)")

        referenced = {n.get(w("val")) for n in document.iter(w("numId"))}
        referenced |= {n.get(w("val")) for n in styles.iter(w("numId"))}
        missing = {r for r in referenced if r and r != "0"} - defined
        if missing:
            problems.append(f"numId without a definition: {sorted(missing)}")

    # Footnotes
    if "word/footnotes.xml" in names:
        footnotes = etree.fromstring(z.read("word/footnotes.xml"))
        defined = {n.get(w("id")) for n in footnotes.findall(w("footnote"))}
        used_ids = {n.get(w("id")) for n in document.iter(w("footnoteReference"))}
        if used_ids - defined:
            problems.append(f"footnote without a definition: {sorted(used_ids - defined)}")

    # Bookmarks paired, anchors resolvable
    starts = {n.get(w("id")) for n in document.iter(w("bookmarkStart"))}
    ends = {n.get(w("id")) for n in document.iter(w("bookmarkEnd"))}
    if starts != ends:
        problems.append("bookmarks not paired")
    bookmarks = {n.get(w("name")) for n in document.iter(w("bookmarkStart"))}
    for link in document.iter(w("hyperlink")):
        anchor = link.get(w("anchor"))
        if anchor and anchor not in bookmarks:
            problems.append(f"missing anchor: {anchor}")

    # No raw line break inside w:t
    for node in document.iter(w("t")):
        if node.text and "\n" in node.text:
            problems.append("w:t contains a line break")
            break

    # Field codes paired
    chars = [n.get(w("fldCharType")) for n in document.iter(w("fldChar"))]
    for name in names:
        if name.startswith("word/footer") or name.startswith("word/header"):
            chars += [
                n.get(w("fldCharType"))
                for n in etree.fromstring(z.read(name)).iter(w("fldChar"))
            ]
    if chars.count("begin") != chars.count("end"):
        problems.append("field codes not paired")

    return problems


@pytest.fixture
def assert_valid():
    def _assert(path) -> None:
        problems = validate_package(Path(path))
        assert not problems, "invalid document:\n" + "\n".join(problems)

    return _assert
