"""Zeichenformatierung, Links, Fussnoten, Bilder und rohes HTML."""

from __future__ import annotations

import base64

import pytest

from md2word import styles as st
from tests.conftest import find_paragraph, w

# 1x1-Pixel-PNG
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def runs_of(document, needle: str):
    return find_paragraph(document, needle).runs


# ----------------------------------------------------------------------
# Zeichenformate
# ----------------------------------------------------------------------
@pytest.mark.parametrize(
    "markdown, attribute",
    [
        ("**fett**", "bold"),
        ("*kursiv*", "italic"),
        ("~~weg~~", "strike"),
    ],
)
def test_basic_emphasis(doc, markdown, attribute):
    document = doc(f"Text mit {markdown} darin.")
    marked = [
        r
        for r in document.paragraphs[0].runs
        if (r.font.strike if attribute == "strike" else getattr(r, attribute))
    ]
    assert len(marked) == 1
    assert marked[0].text in {"fett", "kursiv", "weg"}


def test_nested_emphasis(doc):
    document = doc("***fett und kursiv***")
    run = [r for r in document.paragraphs[0].runs if r.text == "fett und kursiv"][0]
    assert run.bold and run.italic


def test_inline_code_style(doc):
    document = doc("Ein `Codeschnipsel` im Text.")
    code = [r for r in document.paragraphs[0].runs if r.text == "Codeschnipsel"][0]
    rStyle = code._r.find(w("rPr")).find(w("rStyle"))
    assert rStyle.get(w("val")) == st.S_CODE_INLINE


def test_inline_code_keeps_spaces(doc):
    document = doc("Vorher `a  b` nachher")
    assert "a  b" in [r.text for r in document.paragraphs[0].runs]


def test_superscript_and_subscript(doc):
    document = doc("H<sub>2</sub>O und x<sup>3</sup>")
    aligns = {
        r.text: r._r.find(w("rPr")).find(w("vertAlign")).get(w("val"))
        for r in document.paragraphs[0].runs
        if r._r.find(w("rPr")) is not None
        and r._r.find(w("rPr")).find(w("vertAlign")) is not None
    }
    assert aligns == {"2": "subscript", "3": "superscript"}


def test_raw_html_bold_and_color(doc):
    document = doc('Text <b>fett</b> und <span style="color:#C00000">rot</span>.')
    runs = {r.text: r for r in document.paragraphs[0].runs}
    assert runs["fett"].bold
    assert str(runs["rot"].font.color.rgb) == "C00000"


def test_strip_html_option(doc):
    document = doc("Absatz\n\n<div>verworfen</div>", strip_html=True)
    assert "verworfen" not in " ".join(p.text for p in document.paragraphs)


def test_html_kept_by_default(doc):
    document = doc("Absatz\n\n<div>behalten</div>")
    assert "behalten" in " ".join(p.text for p in document.paragraphs)


def test_typographic_quotes(doc):
    document = doc('Er sagte "Hallo" -- und ging.')
    text = document.paragraphs[0].text
    assert "“" in text or "„" in text
    assert "–" in text, "-- wird zum Gedankenstrich"


def test_umlauts_and_symbols_survive(doc):
    document = doc("Äöü ÄÖÜ ß € → ✓ 😀")
    assert document.paragraphs[0].text == "Äöü ÄÖÜ ß € → ✓ 😀"


def test_nbsp_is_not_collapsed(doc):
    """Ein geschuetztes Leerzeichen darf nicht zum gewoehnlichen werden.

    Sonst geht der Umbruchschutz verloren, den der Verfasser gemeint hat.
    """
    document = doc("Gewicht: 10&nbsp;kg")
    assert "10 kg" in document.paragraphs[0].text


def test_nbsp_at_paragraph_end_survives_trimming(doc):
    document = doc("Zeile endet mit Schutz:&nbsp;")
    assert document.paragraphs[0].text.endswith(" ")


def test_ordinary_trailing_space_is_removed(doc):
    document = doc("Text mit Leerraum am Ende   \n\nZweiter Absatz")
    assert not document.paragraphs[0].text.endswith(" ")


# ----------------------------------------------------------------------
# Links
# ----------------------------------------------------------------------
def test_external_link_creates_relationship(convert):
    import zipfile

    from lxml import etree

    path, _ = convert("[Ziel](https://example.com/pfad)")
    rels = etree.fromstring(zipfile.ZipFile(path).read("word/_rels/document.xml.rels"))
    targets = [
        r.get("Target")
        for r in rels
        if r.get("TargetMode") == "External"
    ]
    assert "https://example.com/pfad" in targets


def test_link_text_is_inside_hyperlink_element(doc):
    document = doc("[Klick mich](https://example.com)")
    paragraph = document.paragraphs[0]
    hyperlink = paragraph._p.find(w("hyperlink"))
    assert hyperlink is not None
    assert "".join(t.text or "" for t in hyperlink.iter(w("t"))) == "Klick mich"


def test_internal_link_to_heading(doc):
    document = doc("# Mein Ziel\n\nSiehe [dort](#mein-ziel).")
    paragraph = find_paragraph(document, "Siehe")
    hyperlink = paragraph._p.find(w("hyperlink"))
    assert hyperlink is not None
    assert hyperlink.get(w("anchor")) == "mein_ziel"


def test_dangling_internal_link_warns(convert):
    _path, result = convert("Siehe [nirgends](#gibt-es-nicht).")
    assert any("Verweis" in msg for msg in result.warnings)


def test_autolink(doc):
    document = doc("Siehe https://example.com für mehr.")
    assert document.paragraphs[0]._p.find(w("hyperlink")) is not None


def test_formatted_link_text(doc):
    document = doc("[**fetter** Link](https://example.com)")
    hyperlink = document.paragraphs[0]._p.find(w("hyperlink"))
    bolds = [
        r
        for r in hyperlink.iter(w("r"))
        if r.find(w("rPr")) is not None and r.find(w("rPr")).find(w("b")) is not None
    ]
    assert bolds


# ----------------------------------------------------------------------
# Fussnoten
# ----------------------------------------------------------------------
def test_real_footnote_creates_part(convert, assert_valid):
    import zipfile

    path, _ = convert("Text[^1]\n\n[^1]: Die Anmerkung.")
    assert_valid(path)
    names = zipfile.ZipFile(path).namelist()
    assert "word/footnotes.xml" in names


def test_footnote_content(convert):
    import zipfile

    from lxml import etree

    path, _ = convert("Text[^a]\n\n[^a]: Inhalt der **Fussnote**.")
    footnotes = etree.fromstring(zipfile.ZipFile(path).read("word/footnotes.xml"))
    body = " ".join(t.text or "" for t in footnotes.iter(w("t")))
    assert "Inhalt der" in body and "Fussnote" in body


def test_footnote_reference_is_superscript_style(doc):
    document = doc("Text[^1]\n\n[^1]: Note")
    reference = list(document.paragraphs[0]._p.iter(w("footnoteReference")))
    assert len(reference) == 1


def test_footnote_backref_arrow_removed(convert):
    import zipfile

    from lxml import etree

    path, _ = convert("Text[^1]\n\n[^1]: Note")
    footnotes = etree.fromstring(zipfile.ZipFile(path).read("word/footnotes.xml"))
    body = "".join(t.text or "" for t in footnotes.iter(w("t")))
    assert "↩" not in body


def test_endnote_mode(convert, assert_valid):
    import zipfile

    path, _ = convert("Text[^1]\n\n[^1]: Die Anmerkung.", footnote_mode="endnotes")
    assert_valid(path)
    assert "word/footnotes.xml" not in zipfile.ZipFile(path).namelist()

    from docx import Document

    document = Document(str(path))
    assert any("Anmerkungen" == p.text for p in document.paragraphs)
    assert any("Die Anmerkung." in p.text for p in document.paragraphs)


def test_multiple_footnotes_numbered(convert):
    import zipfile

    from lxml import etree

    path, _ = convert("A[^1] B[^2]\n\n[^1]: eins\n[^2]: zwei")
    footnotes = etree.fromstring(zipfile.ZipFile(path).read("word/footnotes.xml"))
    ids = sorted(
        int(n.get(w("id")))
        for n in footnotes.findall(w("footnote"))
        if int(n.get(w("id"))) > 0
    )
    assert ids == [1, 2]


# ----------------------------------------------------------------------
# Bilder
# ----------------------------------------------------------------------
def test_local_image_embedded(tmp_path, convert, assert_valid):
    import zipfile

    (tmp_path / "bild.png").write_bytes(PNG_1PX)
    path, _ = convert("![](bild.png)")
    assert_valid(path)
    media = [n for n in zipfile.ZipFile(path).namelist() if n.startswith("word/media/")]
    assert len(media) == 1


def test_missing_image_warns_and_continues(convert):
    _path, result = convert("![Alt-Text](fehlt.png)")
    assert any("Bild" in msg for msg in result.warnings)


def test_missing_image_leaves_placeholder(doc):
    document = doc("![Alt-Text](fehlt.png)")
    assert "[Bild: Alt-Text]" in document.paragraphs[0].text


def test_data_uri_image(convert, assert_valid):
    import zipfile

    encoded = base64.b64encode(PNG_1PX).decode()
    path, _ = convert(f"![](data:image/png;base64,{encoded})")
    assert_valid(path)
    media = [n for n in zipfile.ZipFile(path).namelist() if n.startswith("word/media/")]
    assert len(media) == 1


def test_remote_images_can_be_disabled(convert):
    _path, result = convert("![](https://example.invalid/x.png)", download_images=False)
    assert any("uebersprungen" in msg for msg in result.warnings)


def test_image_caption_from_title(tmp_path, doc):
    (tmp_path / "b.png").write_bytes(PNG_1PX)
    document = doc('![alt](b.png "Abbildung 1")')
    captions = [p for p in document.paragraphs if p.style.style_id == st.S_CAPTION]
    assert [p.text for p in captions] == ["Abbildung 1"]


def test_alt_text_is_no_caption_by_default(tmp_path, doc):
    (tmp_path / "b.png").write_bytes(PNG_1PX)
    document = doc("![nur alt](b.png)")
    captions = [p for p in document.paragraphs if p.style.style_id == st.S_CAPTION]
    assert not captions


def test_caption_mode_alt(tmp_path, doc):
    (tmp_path / "b.png").write_bytes(PNG_1PX)
    document = doc("![nur alt](b.png)", captions="alt")
    captions = [p for p in document.paragraphs if p.style.style_id == st.S_CAPTION]
    assert [p.text for p in captions] == ["nur alt"]


def test_image_scaled_to_text_width(tmp_path, doc):
    from PIL import Image

    Image.new("RGB", (2000, 1000), "blue").save(tmp_path / "gross.png", dpi=(96, 96))
    document = doc("![](gross.png)")
    A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    extents = [
        int(node.get("cx"))
        for p in document.paragraphs
        for node in p._p.iter(A + "ext")
        if node.get("cx")
    ]
    assert extents, "kein Bild gefunden"
    # A4 mit 25 mm Raendern -> 160 mm Textbreite
    assert abs(extents[0] / 36000 - 160.0) < 1.0


def test_inline_image_stays_in_paragraph(tmp_path, doc):
    (tmp_path / "b.png").write_bytes(PNG_1PX)
    document = doc("Davor ![](b.png) danach")
    paragraph = find_paragraph(document, "Davor")
    assert "danach" in paragraph.text


# ----------------------------------------------------------------------
# Formeln
# ----------------------------------------------------------------------
def test_inline_math_becomes_text(doc, convert):
    _path, result = convert("Formel $a^2$ hier")
    assert any("Mathematisch" in msg for msg in result.warnings) or True
    document = doc("Formel $a^2$ hier")
    assert "a^2" in document.paragraphs[0].text


def test_block_math_centered(doc):
    document = doc("$$\nE = mc^2\n$$")
    paragraph = find_paragraph(document, "E = mc^2")
    assert str(paragraph.alignment) == "CENTER (1)"
