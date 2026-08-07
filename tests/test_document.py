"""Dokumentrahmen: Front Matter, Titelseite, Verzeichnis, Kopf-/Fusszeile, Themes."""

from __future__ import annotations

import zipfile

import pytest
from docx import Document
from docx.shared import Mm
from lxml import etree

from md2word import styles as st
from md2word.config import PAGE_SIZES, THEMES, Config
from tests.conftest import find_paragraph, texts, w


def instr_texts(document) -> list[str]:
    return [
        (n.text or "").strip()
        for n in document.element.body.iter(w("instrText"))
    ]


# ----------------------------------------------------------------------
# Front Matter
# ----------------------------------------------------------------------
def test_front_matter_sets_core_properties(convert):
    path, _ = convert(
        "---\n"
        "title: Mein Titel\n"
        "author: Jane Doe\n"
        "keywords: a, b\n"
        "subject: Ein Thema\n"
        "---\n\n"
        "# Inhalt\n"
    )
    props = Document(str(path)).core_properties
    assert props.title == "Mein Titel"
    assert props.author == "Jane Doe"
    assert props.keywords == "a, b"
    assert props.subject == "Ein Thema"


def test_front_matter_is_not_rendered(doc):
    document = doc("---\ntitle: Weg damit\n---\n\nNur dieser Text.")
    assert "title:" not in " ".join(texts(document))


def test_front_matter_enables_toc(doc):
    document = doc("---\ntoc: true\n---\n\n# Kapitel")
    assert any("TOC" in t for t in instr_texts(document))


def test_front_matter_list_author(convert):
    path, _ = convert("---\nauthor:\n  - Erste Person\n  - Zweite Person\n---\n\nText")
    assert Document(str(path)).core_properties.author == "Erste Person, Zweite Person"


def test_front_matter_date_object(convert):
    path, _ = convert("---\ntitle: T\ndate: 2026-08-07\n---\n\nText", title_page=True)
    assert any("2026-08-07" in t for t in texts(Document(str(path))))


def test_cli_option_beats_front_matter():
    """Was explizit auf der Kommandozeile steht, gewinnt gegen das Front Matter."""
    from md2word.converter import _apply_front_matter

    config = Config(theme="modern")
    config._explicit = {"theme"}
    merged = _apply_front_matter(config, {"theme": "classic"})
    assert merged.theme == "modern"

    config2 = Config(theme="modern")
    merged2 = _apply_front_matter(config2, {"theme": "classic"})
    assert merged2.theme == "classic"


def test_broken_front_matter_is_tolerated(doc):
    document = doc("---\n: : nicht: gueltig: yaml\n---\n\nDer Text kommt trotzdem.")
    assert any("Der Text kommt trotzdem." in t for t in texts(document))


def test_unknown_front_matter_keys_are_kept(convert):
    from md2word.converter import _apply_front_matter

    merged = _apply_front_matter(Config(), {"projektnummer": 4711})
    assert merged._extra["projektnummer"] == 4711


# ----------------------------------------------------------------------
# Titelseite
# ----------------------------------------------------------------------
def test_title_page_layout(doc):
    document = doc(
        "---\ntitle: Der Titel\nsubtitle: Der Untertitel\nauthor: Wer\n---\n\n# Kapitel",
        title_page=True,
    )
    styles = [p.style.style_id for p in document.paragraphs]
    assert "Title" in styles and "Subtitle" in styles
    assert styles.index("Title") < styles.index("Heading1")
    breaks = [n for n in document.element.body.iter(w("br")) if n.get(w("type")) == "page"]
    assert breaks, "nach der Titelseite folgt ein Seitenumbruch"


def test_title_without_title_page_is_inline(doc):
    document = doc("---\ntitle: Der Titel\n---\n\nNur Text, keine Ueberschrift.")
    assert find_paragraph(document, "Der Titel").style.style_id == "Title"


def test_title_skipped_when_document_starts_with_h1(doc):
    document = doc("---\ntitle: Meta-Titel\n---\n\n# Eigene Ueberschrift")
    assert "Title" not in [p.style.style_id for p in document.paragraphs]


# ----------------------------------------------------------------------
# Inhaltsverzeichnis
# ----------------------------------------------------------------------
def test_toc_field_present(doc):
    document = doc("# Eins\n\n## Zwei", toc=True)
    fields = instr_texts(document)
    assert any(f.startswith("TOC") for f in fields)


def test_toc_depth_in_field(doc):
    document = doc("# Eins", toc=True, toc_depth=2)
    assert any('"1-2"' in f for f in instr_texts(document))


def test_toc_heading_uses_configured_title(doc):
    document = doc("# Eins", toc=True, toc_title="Übersicht")
    assert find_paragraph(document, "Übersicht").style.style_id == "TOCHeading"


def test_toc_triggers_field_update(convert):
    path, _ = convert("# Eins", toc=True)
    settings = etree.fromstring(zipfile.ZipFile(path).read("word/settings.xml"))
    node = settings.find(w("updateFields"))
    assert node is not None and node.get(w("val")) == "true"


# ----------------------------------------------------------------------
# Nummerierte Ueberschriften
# ----------------------------------------------------------------------
def test_numbered_headings_attach_numbering_to_styles(convert, assert_valid):
    path, _ = convert("# Eins\n\n## Eins-Eins", number_headings=True)
    assert_valid(path)
    styles = etree.fromstring(zipfile.ZipFile(path).read("word/styles.xml"))
    heading1 = [
        s for s in styles.findall(w("style")) if s.get(w("styleId")) == "Heading1"
    ][0]
    numPr = heading1.find(w("pPr")).find(w("numPr"))
    assert numPr is not None, "Heading 1 muss mit der Nummerierung verknuepft sein"


# ----------------------------------------------------------------------
# Kopf- und Fusszeile
# ----------------------------------------------------------------------
def test_page_numbers_create_footer_field(convert):
    path, _ = convert("Text", page_numbers=True)
    z = zipfile.ZipFile(path)
    footers = [n for n in z.namelist() if n.startswith("word/footer")]
    assert footers
    content = etree.fromstring(z.read(footers[0]))
    instructions = [(n.text or "").strip() for n in content.iter(w("instrText"))]
    assert "PAGE" in instructions and "NUMPAGES" in instructions


def test_header_text(convert):
    path, _ = convert("Text", header_text="Vertraulich")
    z = zipfile.ZipFile(path)
    headers = [n for n in z.namelist() if n.startswith("word/header")]
    assert headers
    body = "".join(
        t.text or "" for t in etree.fromstring(z.read(headers[0])).iter(w("t"))
    )
    assert "Vertraulich" in body


def test_footer_text_and_page_numbers_combined(convert, assert_valid):
    path, _ = convert("Text", page_numbers=True, footer_text="Firma AG")
    assert_valid(path)
    z = zipfile.ZipFile(path)
    footer = [n for n in z.namelist() if n.startswith("word/footer")][0]
    tree = etree.fromstring(z.read(footer))
    assert "Firma AG" in "".join(t.text or "" for t in tree.iter(w("t")))
    assert "PAGE" in [(n.text or "").strip() for n in tree.iter(w("instrText"))]


# ----------------------------------------------------------------------
# Seitenlayout
# ----------------------------------------------------------------------
@pytest.mark.parametrize("size", sorted(PAGE_SIZES))
def test_page_sizes(doc, size):
    document = doc("Text", page_size=size)
    section = document.sections[0]
    expected_w, expected_h = PAGE_SIZES[size]
    assert abs(section.page_width.mm - expected_w) < 0.5
    assert abs(section.page_height.mm - expected_h) < 0.5


def test_landscape_swaps_dimensions(doc):
    document = doc("Text", page_size="a4", landscape=True)
    section = document.sections[0]
    assert section.page_width > section.page_height


def test_margins(doc):
    document = doc("Text", margin_top=10, margin_bottom=15, margin_left=20, margin_right=25)
    section = document.sections[0]
    assert abs(section.top_margin.mm - 10) < 0.5
    assert abs(section.left_margin.mm - 20) < 0.5
    assert abs(section.right_margin.mm - 25) < 0.5


# ----------------------------------------------------------------------
# Themes und Typografie
# ----------------------------------------------------------------------
@pytest.mark.parametrize("theme", sorted(THEMES))
def test_themes_apply_fonts(doc, theme):
    document = doc("# Titel\n\nText", theme=theme)
    normal = [s for s in document.styles if s.style_id == "Normal"][0]
    assert normal.font.name == THEMES[theme]["body_font"]


def test_font_size_and_spacing(doc):
    document = doc("Text", font_size=14, line_spacing=2.0)
    normal = [s for s in document.styles if s.style_id == "Normal"][0]
    assert normal.font.size.pt == 14
    assert normal.paragraph_format.line_spacing == 2.0


def test_custom_fonts_override_theme(doc):
    document = doc("# T\n\nText", theme="modern", body_font="Georgia", code_font="Menlo")
    normal = [s for s in document.styles if s.style_id == "Normal"][0]
    assert normal.font.name == "Georgia"


def test_document_language_is_set(convert):
    path, _ = convert("Text", lang="de-AT")
    styles = etree.fromstring(zipfile.ZipFile(path).read("word/styles.xml"))
    langs = [n.get(w("val")) for n in styles.iter(w("lang"))]
    assert "de-AT" in langs


def test_german_quotes(doc):
    document = doc('Er sagte "Hallo".', lang="de-DE")
    assert "„Hallo“" in document.paragraphs[0].text


def test_english_quotes(doc):
    document = doc('He said "Hello".', lang="en-US")
    assert "“Hello”" in document.paragraphs[0].text


@pytest.mark.parametrize(
    "lang, opening, closing",
    [
        ("de-DE", "„", "“"),
        ("de-AT", "„", "“"),
        ("pl-PL", "„", "“"),
        ("fr-FR", "«", "»"),
        ("es-ES", "«", "»"),
        ("it-IT", "«", "»"),
        ("ru-RU", "«", "»"),
        ("en-US", "“", "”"),
        ("nl-NL", "“", "”"),
        ("", "“", "”"),
    ],
)
def test_quotes_per_language(doc, lang, opening, closing):
    document = doc('Er sagte "Wort" dazu.', lang=lang or "en")
    text = document.paragraphs[0].text
    assert opening in text and closing in text, f"{lang}: {text!r}"
    assert '"' not in text, f"{lang}: gerades Anfuehrungszeichen uebrig: {text!r}"


def test_every_quote_set_has_exactly_four_entries():
    """markdown-it greift ueber den Index zu.

    Ein laengerer String liefert stillschweigend die falschen Zeichen -
    genau so war das schliessende Guillemet einmal ein Leerzeichen.
    """
    from md2word.parser import _QUOTES, quotes_for

    for name, quotes in _QUOTES.items():
        assert len(quotes) == 4, f"{name}: {len(quotes)} statt 4 Eintraege"

    for lang in ("de", "fr", "en", "ru", "xx", ""):
        assert len(quotes_for(lang)) == 4


def test_french_quotes_use_narrow_space(doc):
    """Franzoesische Typografie setzt Guillemets mit schmalem Abstand."""
    document = doc('Il a dit "Bonjour" ensuite.', lang="fr-FR")
    assert "« Bonjour »" in document.paragraphs[0].text


# ----------------------------------------------------------------------
# Referenzdokument
# ----------------------------------------------------------------------
def test_reference_doc_keeps_its_styles(tmp_path, convert):
    reference = tmp_path / "vorlage.docx"
    template = Document()
    template.styles["Normal"].font.name = "Garamond"
    template.styles["Normal"].font.size = Mm(0).__class__(190500)  # 15 pt
    template.add_paragraph("Dieser Inhalt muss verschwinden.")
    template.save(str(reference))

    path, _ = convert("# Neu\n\nText", reference_doc=str(reference))
    document = Document(str(path))
    assert document.styles["Normal"].font.name == "Garamond"
    assert "Dieser Inhalt muss verschwinden." not in " ".join(texts(document))
    assert "Neu" in " ".join(texts(document))


def test_reference_doc_gets_missing_styles(tmp_path, convert, assert_valid):
    reference = tmp_path / "vorlage.docx"
    Document().save(str(reference))

    path, _ = convert("```python\nx = 1\n```", reference_doc=str(reference))
    assert_valid(path)
    document = Document(str(path))
    assert find_paragraph(document, "x = 1").style.style_id == st.S_CODE_BLOCK


def test_missing_reference_doc_raises(convert):
    with pytest.raises(FileNotFoundError):
        convert("Text", reference_doc="/gibt/es/nicht.docx")
