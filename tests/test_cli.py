"""Kommandozeile: Argumente, Dateinamen, Stapelverarbeitung, Fehlerfaelle."""

from __future__ import annotations

import io
import os
import sys

import pytest
from docx import Document

from md2word.cli import main


@pytest.fixture
def workdir(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    return tmp_path


def write(path, text="# Titel\n\nEin Absatz.\n"):
    path.write_text(text, encoding="utf-8")
    return path


# ----------------------------------------------------------------------
def test_default_output_name(workdir, capsys):
    write(workdir / "bericht.md")
    assert main(["bericht.md"]) == 0
    assert (workdir / "bericht.docx").is_file()


def test_explicit_output(workdir):
    write(workdir / "a.md")
    assert main(["a.md", "-o", "ziel.docx"]) == 0
    assert (workdir / "ziel.docx").is_file()


def test_output_directory_is_created(workdir):
    write(workdir / "a.md")
    assert main(["a.md", "-d", "build/unterordner"]) == 0
    assert (workdir / "build" / "unterordner" / "a.docx").is_file()


def test_batch_conversion(workdir):
    for name in ("eins.md", "zwei.md", "drei.md"):
        write(workdir / name)
    assert main(["eins.md", "zwei.md", "drei.md", "-d", "out"]) == 0
    assert sorted(p.name for p in (workdir / "out").iterdir()) == [
        "drei.docx",
        "eins.docx",
        "zwei.docx",
    ]


def test_glob_pattern(workdir):
    for name in ("a.md", "b.md"):
        write(workdir / name)
    assert main(["*.md", "-d", "out"]) == 0
    assert len(list((workdir / "out").iterdir())) == 2


def test_existing_target_is_protected(workdir, capsys):
    write(workdir / "a.md")
    (workdir / "a.docx").write_text("alt")
    assert main(["a.md"]) == 1
    assert (workdir / "a.docx").read_text() == "alt"
    assert "existiert bereits" in capsys.readouterr().err


def test_force_overwrites(workdir):
    write(workdir / "a.md")
    (workdir / "a.docx").write_text("alt")
    assert main(["a.md", "--force"]) == 0
    assert (workdir / "a.docx").read_bytes()[:2] == b"PK"


def test_missing_input_reports_error(workdir, capsys):
    assert main(["fehlt.md"]) == 1
    assert "nicht gefunden" in capsys.readouterr().err


def test_output_with_multiple_inputs_is_rejected(workdir, capsys):
    write(workdir / "a.md")
    write(workdir / "b.md")
    with pytest.raises(SystemExit) as exc:
        main(["a.md", "b.md", "-o", "eins.docx"])
    assert exc.value.code == 2


def test_stdin_input(workdir, monkeypatch):
    monkeypatch.setattr(sys, "stdin", io.StringIO("# Aus der Pipe\n\nText"))
    assert main(["-", "-o", "pipe.docx"]) == 0
    document = Document(str(workdir / "pipe.docx"))
    assert any("Aus der Pipe" in p.text for p in document.paragraphs)


def test_quiet_suppresses_status(workdir, capsys):
    write(workdir / "a.md")
    main(["a.md", "--quiet"])
    assert capsys.readouterr().out == ""


def test_status_line_mentions_target(workdir, capsys):
    write(workdir / "a.md")
    main(["a.md"])
    assert "a.docx" in capsys.readouterr().out


def test_no_arguments_shows_help(capsys):
    assert main([]) == 2
    assert "usage" in capsys.readouterr().out.lower()


def test_version_flag(capsys):
    with pytest.raises(SystemExit) as exc:
        main(["--version"])
    assert exc.value.code == 0
    assert "md2word" in capsys.readouterr().out


def test_list_pygments_styles(capsys):
    assert main(["--list-pygments-styles"]) == 0
    assert "friendly" in capsys.readouterr().out


def test_batch_continues_after_failure(workdir, capsys):
    write(workdir / "gut.md")
    assert main(["fehlt.md", "gut.md", "-d", "out"]) == 1
    assert (workdir / "out" / "gut.docx").is_file()


# ----------------------------------------------------------------------
# Argumentauswertung
# ----------------------------------------------------------------------
def test_margin_sets_all_sides(workdir):
    write(workdir / "a.md")
    main(["a.md", "--margin", "12"])
    section = Document(str(workdir / "a.docx")).sections[0]
    for value in (section.top_margin, section.bottom_margin, section.left_margin):
        assert abs(value.mm - 12) < 0.5


def test_specific_margin_wins_over_general(workdir):
    write(workdir / "a.md")
    main(["a.md", "--margin", "12", "--margin-left", "40"])
    section = Document(str(workdir / "a.docx")).sections[0]
    assert abs(section.left_margin.mm - 40) < 0.5
    assert abs(section.top_margin.mm - 12) < 0.5


def test_explicit_flags_are_detected():
    from md2word.cli import _explicit_options, build_parser

    parser = build_parser()
    seen = _explicit_options(["a.md", "--theme", "modern", "--toc"], parser)
    assert "theme" in seen and "toc" in seen
    assert "page_size" not in seen


def test_no_highlight_maps_to_config():
    from md2word.cli import _explicit_options, build_parser, config_from_args

    parser = build_parser()
    argv = ["a.md", "--no-highlight"]
    args = parser.parse_args(argv)
    config = config_from_args(args, _explicit_options(argv, parser))
    assert config.highlight is False


def test_cli_metadata_overrides_front_matter(workdir):
    write(workdir / "a.md", "---\ntitle: Aus dem Front Matter\n---\n\nText")
    main(["a.md", "--title", "Von der Kommandozeile"])
    assert Document(str(workdir / "a.docx")).core_properties.title == "Von der Kommandozeile"


def test_pygments_style_derives_code_background():
    from md2word.cli import _explicit_options, build_parser, config_from_args

    parser = build_parser()
    argv = ["a.md", "--pygments-style", "monokai"]
    args = parser.parse_args(argv)
    config = config_from_args(args, _explicit_options(argv, parser))
    assert config.code_bg != "F5F5F5", "dunkles Schema braucht dunklen Hintergrund"
