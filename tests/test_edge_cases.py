"""Randfaelle: leere Eingaben, kaputtes Markup, tiefe Verschachtelung, Grossdokumente."""

from __future__ import annotations

import pytest

from md2word import styles as st
from tests.conftest import find_paragraph, texts, w


# ----------------------------------------------------------------------
# Leere und minimale Eingaben
# ----------------------------------------------------------------------
@pytest.mark.parametrize("source", ["", "   ", "\n\n\n", "\t\n \n"])
def test_empty_input_produces_valid_document(convert, assert_valid, source):
    path, _ = convert(source)
    assert_valid(path)


def test_only_front_matter(convert, assert_valid):
    path, _ = convert("---\ntitle: Nur Metadaten\n---\n")
    assert_valid(path)


def test_single_character(doc):
    assert doc("x").paragraphs[0].text == "x"


def test_bom_is_stripped(tmp_path):
    from md2word.config import Config
    from md2word.converter import convert_file

    source = tmp_path / "bom.md"
    source.write_bytes("﻿# Titel\n\nText".encode("utf-8"))
    target = tmp_path / "bom.docx"
    convert_file(str(source), str(target), Config())

    from docx import Document

    document = Document(str(target))
    assert document.paragraphs[0].text == "Titel"
    assert not document.paragraphs[0].text.startswith("﻿")


@pytest.mark.parametrize("newline", ["\r\n", "\r"])
def test_windows_and_mac_line_endings(doc, newline):
    document = doc(f"# Titel{newline}{newline}Ein Absatz.{newline}")
    assert [t for t in texts(document) if t] == ["Titel", "Ein Absatz."]


# ----------------------------------------------------------------------
# Kaputtes oder ungewoehnliches Markup
# ----------------------------------------------------------------------
@pytest.mark.parametrize(
    "source",
    [
        "**nicht geschlossen",
        "[Link ohne Ziel",
        "| A | B\n|---",
        "```\nnie geschlossen",
        "<div><span>nicht geschlossen",
        "> Zitat\n>> ohne Leerzeile",
        "![Bild ohne Klammer](",
        "- Punkt\n    - stark eingerueckt\n\t- Tabulator",
        "#" * 12 + " Zu viele Rauten",
        "|||\n|-|-|\n|||",
    ],
)
def test_malformed_markdown_does_not_crash(convert, assert_valid, source):
    path, _ = convert(source)
    assert_valid(path)


def test_unclosed_html_is_tolerated(convert, assert_valid):
    path, _ = convert("Text <b>fett ohne Ende\n\nNaechster Absatz")
    assert_valid(path)


def test_script_and_style_are_dropped(doc):
    document = doc("Text\n\n<script>alert(1)</script>\n\n<style>p{color:red}</style>")
    body = " ".join(texts(document))
    assert "alert" not in body and "color:red" not in body


def test_html_table_is_rendered(doc):
    document = doc("<table><tr><td>Zelle A</td><td>Zelle B</td></tr></table>")
    assert document.tables, "auch rohes HTML-Tabellenmarkup soll eine Tabelle ergeben"
    assert document.tables[0].rows[0].cells[0].text == "Zelle A"


# ----------------------------------------------------------------------
# Tiefe Verschachtelung
# ----------------------------------------------------------------------
def test_deeply_nested_lists(convert, assert_valid):
    source = "\n".join("  " * level + "- Ebene " + str(level) for level in range(12))
    path, _ = convert(source)
    assert_valid(path)


def test_list_levels_are_capped_at_nine(doc):
    source = "\n".join("  " * level + "- E" + str(level) for level in range(12))
    document = doc(source)
    levels = []
    for paragraph in document.paragraphs:
        pPr = paragraph._p.find(w("pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(w("numPr"))
        if numPr is not None:
            levels.append(int(numPr.find(w("ilvl")).get(w("val"))))
    assert levels, "Listenabsaetze erwartet"
    assert max(levels) <= 8, f"Word kennt nur die Ebenen 0-8, gefunden: {max(levels)}"


def test_deeply_nested_quotes(convert, assert_valid):
    path, _ = convert("> " * 15 + "Sehr tief verschachtelt")
    assert_valid(path)


def test_table_inside_list(convert, assert_valid):
    path, _ = convert("- Punkt\n\n  | A | B |\n  |---|---|\n  | 1 | 2 |\n")
    assert_valid(path)
    from docx import Document

    assert Document(str(path)).tables


def test_list_inside_table_cell(convert, assert_valid):
    path, _ = convert("| Spalte |\n|---|\n| <ul><li>eins</li><li>zwei</li></ul> |")
    assert_valid(path)


def test_quote_containing_list_and_code(convert, assert_valid):
    path, _ = convert("> - Punkt\n> - Anderer\n>\n> ```\n> code\n> ```")
    assert_valid(path)


def test_mixed_nesting(convert, assert_valid):
    source = (
        "1. Erst\n"
        "   > Zitat im Listenpunkt\n"
        "   >\n"
        "   > - Liste im Zitat im Listenpunkt\n"
        "2. Zweitens\n"
        "   ```python\n"
        "   x = 1\n"
        "   ```\n"
    )
    path, _ = convert(source)
    assert_valid(path)


# ----------------------------------------------------------------------
# Sonderzeichen und Kodierung
# ----------------------------------------------------------------------
def test_xml_special_characters(doc):
    document = doc("Zeichen: < > & \" ' und ]]>")
    text = document.paragraphs[0].text
    assert "<" in text and ">" in text and "&" in text


def test_control_characters_are_survivable(convert, assert_valid):
    path, _ = convert("Text mit Tabulator\tund Formfeed\x0c dazwischen")
    assert_valid(path)


def test_emoji_and_cjk(doc):
    document = doc("Emoji 🎉 und CJK 日本語 und Arabisch مرحبا")
    assert "🎉" in document.paragraphs[0].text
    assert "日本語" in document.paragraphs[0].text


def test_very_long_line(convert, assert_valid):
    path, _ = convert("Wort " * 5000)
    assert_valid(path)


def test_many_headings(convert, assert_valid):
    source = "\n\n".join(f"## Abschnitt {i}" for i in range(300))
    path, _ = convert(source, toc=True)
    assert_valid(path)


def test_duplicate_heading_anchors_are_unique(doc):
    document = doc("# Gleich\n\n# Gleich\n\n# Gleich")
    names = [n.get(w("name")) for n in document.element.body.iter(w("bookmarkStart"))]
    assert len(names) == len(set(names)), f"Anker mehrfach vergeben: {names}"


def test_long_heading_bookmark_is_truncated(doc):
    document = doc("# " + "Sehr langer Titel " * 10)
    names = [n.get(w("name")) for n in document.element.body.iter(w("bookmarkStart"))]
    assert all(len(n) <= 40 for n in names), "Word erlaubt hoechstens 40 Zeichen"


def test_heading_with_only_symbols(convert, assert_valid):
    path, _ = convert("# !!! ??? ***\n\nText")
    assert_valid(path)


# ----------------------------------------------------------------------
# Tabellen-Randfaelle
# ----------------------------------------------------------------------
def test_table_with_empty_cells(convert, assert_valid):
    path, _ = convert("| A | B |\n|---|---|\n|   | 2 |\n| 3 |   |")
    assert_valid(path)


def test_table_with_pipes_in_code(doc):
    document = doc("| Spalte |\n|---|\n| `a \\| b` |")
    assert "|" in document.tables[0].rows[1].cells[0].text


def test_wide_table_stays_within_page(doc):
    header = "| " + " | ".join(f"S{i}" for i in range(12)) + " |"
    separator = "|" + "---|" * 12
    row = "| " + " | ".join(str(i) for i in range(12)) + " |"
    document = doc(f"{header}\n{separator}\n{row}")
    total = sum(cell.width.mm for cell in document.tables[0].rows[0].cells)
    assert total < 165, f"Tabelle laeuft ueber den Satzspiegel: {total} mm"


def test_two_tables_in_a_row(convert, assert_valid):
    path, _ = convert("| A |\n|---|\n| 1 |\n\n| B |\n|---|\n| 2 |")
    assert_valid(path)
    from docx import Document

    assert len(Document(str(path)).tables) == 2


# ----------------------------------------------------------------------
# Fussnoten-Randfaelle
# ----------------------------------------------------------------------
def test_footnote_without_definition(convert, assert_valid):
    path, _ = convert("Text mit [^unbekannt] Verweis.")
    assert_valid(path)


def test_footnote_with_multiple_paragraphs(convert, assert_valid):
    path, _ = convert("Text[^1]\n\n[^1]: Erster Absatz.\n\n    Zweiter Absatz.")
    assert_valid(path)


def test_footnote_with_link(convert, assert_valid):
    path, _ = convert("Text[^1]\n\n[^1]: Siehe [Quelle](https://example.com).")
    assert_valid(path)


def test_footnote_referenced_twice(convert, assert_valid):
    path, _ = convert("A[^n] und B[^n]\n\n[^n]: Die Note")
    assert_valid(path)


# ----------------------------------------------------------------------
# Wiederholbarkeit
# ----------------------------------------------------------------------
def test_conversion_is_deterministic(tmp_path):
    """Zweimal dasselbe Markdown ergibt dasselbe document.xml."""
    import zipfile

    from md2word.config import Config
    from md2word.converter import convert_text

    source = "# Titel\n\n- eins\n- zwei\n\n[Link](https://example.com)\n\nText[^1]\n\n[^1]: Note"
    outputs = []
    for index in range(2):
        target = tmp_path / f"d{index}.docx"
        convert_text(source, str(target), Config())
        outputs.append(zipfile.ZipFile(target).read("word/document.xml"))
    assert outputs[0] == outputs[1]
