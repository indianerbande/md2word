"""Language handling: the document's own strings follow --lang."""

from __future__ import annotations

import pytest

from md2word import i18n
from md2word.config import Config
from tests.conftest import find_paragraph, texts


# ----------------------------------------------------------------------
# The lookup itself
# ----------------------------------------------------------------------
def test_default_language_is_english():
    assert Config().lang == "en-US"


def test_primary_subtag_is_extracted():
    assert i18n.language_of("de-AT") == "de"
    assert i18n.language_of("de_DE") == "de"
    assert i18n.language_of("EN") == "en"
    assert i18n.language_of("") == "en"


def test_unknown_language_falls_back_to_english():
    assert i18n.translate("xx-XX", "toc_title") == "Table of Contents"
    assert i18n.translate("", "endnotes_title") == "Notes"


def test_every_language_has_every_key():
    """A partial translation is allowed, but a typo in a key is not."""
    english = set(i18n._STRINGS["en"])
    for name, table in i18n._STRINGS.items():
        unknown = set(table) - english
        assert not unknown, f"{name} has keys English does not: {sorted(unknown)}"


def test_placeholders_survive_translation():
    for lang in i18n.available_languages():
        rendered = i18n.translate(lang, "image_placeholder", label="x.png")
        assert "x.png" in rendered
        assert "{" not in rendered
        note = i18n.translate(lang, "generated_note", source="a.md")
        assert "a.md" in note and "{" not in note


# ----------------------------------------------------------------------
# Effect on the generated document
# ----------------------------------------------------------------------
@pytest.mark.parametrize(
    "lang, expected",
    [
        ("en-US", "Table of Contents"),
        ("de-DE", "Inhaltsverzeichnis"),
        ("de-AT", "Inhaltsverzeichnis"),
        ("fr-FR", "Table des matières"),
        ("es-ES", "Índice"),
        ("it-IT", "Indice"),
        ("sv-SE", "Table of Contents"),  # not translated -> English
    ],
)
def test_toc_title_follows_language(doc, lang, expected):
    document = doc("# Chapter", toc=True, lang=lang)
    assert expected in texts(document)


def test_explicit_toc_title_wins_over_language(doc):
    document = doc("# Chapter", toc=True, lang="de-DE", toc_title="Übersicht")
    assert "Übersicht" in texts(document)
    assert "Inhaltsverzeichnis" not in texts(document)


@pytest.mark.parametrize(
    "lang, expected",
    [("en-US", "Notes"), ("de-DE", "Anmerkungen"), ("fr-FR", "Notes")],
)
def test_endnotes_heading_follows_language(doc, lang, expected):
    document = doc("Text[^1]\n\n[^1]: A note.", footnote_mode="endnotes", lang=lang)
    assert any(p.text == expected for p in document.paragraphs)


@pytest.mark.parametrize(
    "lang, expected",
    [("en-US", "[Image: alt]"), ("de-DE", "[Bild: alt]"), ("es-ES", "[Imagen: alt]")],
)
def test_image_placeholder_follows_language(doc, lang, expected):
    document = doc("![alt](missing.png)", lang=lang)
    assert expected in document.paragraphs[0].text


def test_toc_placeholder_is_translated(doc):
    english = doc("# Chapter", toc=True, lang="en-US")
    german = doc("# Chapter", toc=True, lang="de-DE")
    assert any("Update Field" in t for t in texts(english))
    assert any("Felder aktualisieren" in t for t in texts(german))


@pytest.mark.parametrize(
    "lang, expected",
    [("en-US", "Generated with md2word"), ("de-DE", "Erzeugt mit md2word")],
)
def test_generated_note_in_core_properties(tmp_path, lang, expected):
    """The note only exists when converting a named file, not raw text."""
    from docx import Document

    from md2word.converter import convert_file

    source = tmp_path / "source.md"
    source.write_text("# Title\n\nText", encoding="utf-8")
    target = tmp_path / "out.docx"
    convert_file(str(source), str(target), Config(lang=lang))

    comments = Document(str(target)).core_properties.comments or ""
    assert expected in comments
    assert "source.md" in comments


def test_title_page_date_is_iso(doc):
    """Without an explicit date, ISO 8601 is used - unambiguous everywhere."""
    import datetime
    import re

    document = doc("---\ntitle: T\n---\n\nText", title_page=True)
    today = datetime.date.today().isoformat()
    assert any(re.fullmatch(r"\d{4}-\d{2}-\d{2}", t) and t == today for t in texts(document))


def test_untitled_placeholder_follows_language(doc):
    english = doc("Text", title_page=True)
    german = doc("Text", title_page=True, lang="de-DE")
    assert "Untitled" in texts(english)
    assert "Ohne Titel" in texts(german)


# ----------------------------------------------------------------------
# Interaction with quotation marks (parser side)
# ----------------------------------------------------------------------
def test_default_language_yields_english_quotes(doc):
    document = doc('He said "Hello".')
    assert "“Hello”" in document.paragraphs[0].text


def test_front_matter_language_reaches_both_layers(doc):
    """lang from the front matter must drive quotes *and* the TOC title."""
    document = doc('---\nlang: de-DE\n---\n\n# Chapter\n\nShe said "Hello".', toc=True)
    body = texts(document)
    assert "Inhaltsverzeichnis" in body
    assert any("„Hello“" in t for t in body)
