"""Blockelemente: Ueberschriften, Listen, Zitate, Code, Tabellen, Trennlinien."""

from __future__ import annotations

import pytest
from docx.shared import Mm

from md2word import styles as st
from tests.conftest import find_paragraph, numbering_of, paragraph_styles, texts, w


# ----------------------------------------------------------------------
# Ueberschriften
# ----------------------------------------------------------------------
def test_headings_map_to_word_levels(doc):
    document = doc("# Eins\n\n## Zwei\n\n### Drei\n\n#### Vier\n\n##### Fuenf\n\n###### Sechs")
    styles = [s for s in paragraph_styles(document) if s.startswith("Heading")]
    assert styles == [f"Heading{i}" for i in range(1, 7)]


def test_heading_keeps_inline_formatting(doc):
    document = doc("# Ein **fetter** Titel")
    heading = find_paragraph(document, "fetter")
    assert [(r.text, bool(r.bold)) for r in heading.runs] == [
        ("Ein ", False),
        ("fetter", True),
        (" Titel", False),
    ]


def test_heading_gets_bookmark(doc):
    document = doc("# Mein Abschnitt\n\nText")
    names = [n.get(w("name")) for n in document.element.body.iter(w("bookmarkStart"))]
    assert "mein_abschnitt" in names, f"Word erlaubt keine Bindestriche: {names}"


def test_break_on_h1_inserts_page_break(doc):
    document = doc("# Erstes\n\nText\n\n# Zweites", break_on_h1=True)
    breaks = [
        n
        for n in document.element.body.iter(w("br"))
        if n.get(w("type")) == "page"
    ]
    assert len(breaks) == 1, "genau ein Umbruch vor der zweiten H1"


# ----------------------------------------------------------------------
# Absaetze und Zeilenumbrueche
# ----------------------------------------------------------------------
def test_hard_line_break_becomes_br(doc):
    document = doc("Zeile eins  \nZeile zwei")
    paragraph = document.paragraphs[0]
    assert len(paragraph._p.findall(".//" + w("br"))) == 1
    assert "Zeile eins" in paragraph.text and "Zeile zwei" in paragraph.text


def test_soft_break_stays_one_paragraph(doc):
    document = doc("Zeile eins\nZeile zwei")
    assert len([p for p in document.paragraphs if p.text.strip()]) == 1


def test_paragraphs_are_separate(doc):
    document = doc("Erster Absatz.\n\nZweiter Absatz.")
    assert [t for t in texts(document) if t] == ["Erster Absatz.", "Zweiter Absatz."]


# ----------------------------------------------------------------------
# Listen
# ----------------------------------------------------------------------
def test_bullet_list_uses_numbering(doc):
    document = doc("- eins\n- zwei\n- drei")
    items = [p for p in document.paragraphs if p.text in {"eins", "zwei", "drei"}]
    assert len(items) == 3
    numbers = {numbering_of(p) for p in items}
    assert len(numbers) == 1, "alle Punkte teilen dieselbe Listendefinition"
    assert numbers.pop()[1] == 0


def test_nested_list_levels(doc):
    document = doc("- a\n  - b\n    - c")
    levels = {
        p.text: numbering_of(p)[1]
        for p in document.paragraphs
        if p.text in {"a", "b", "c"}
    }
    assert levels == {"a": 0, "b": 1, "c": 2}


def test_ordered_lists_restart(doc):
    document = doc("1. eins\n2. zwei\n\nText dazwischen\n\n1. neu\n2. wieder")
    first = numbering_of(find_paragraph(document, "eins"))[0]
    second = numbering_of(find_paragraph(document, "neu"))[0]
    assert first != second, "zwei Listen brauchen eigene numIds, sonst zaehlt Word weiter"


def test_ordered_list_start_attribute(convert):
    path, _ = convert("5. fuenf\n6. sechs")
    from docx import Document
    from lxml import etree
    import zipfile

    numbering = etree.fromstring(zipfile.ZipFile(path).read("word/numbering.xml"))
    overrides = [n.get(w("val")) for n in numbering.iter(w("startOverride"))]
    assert "5" in overrides


def test_task_list_checkboxes(doc):
    document = doc("- [x] erledigt\n- [ ] offen")
    assert find_paragraph(document, "erledigt").text.startswith("☒")
    assert find_paragraph(document, "offen").text.startswith("☐")


def test_list_item_with_second_paragraph(doc):
    document = doc("- Punkt\n\n  Zweiter Absatz des Punktes\n\n- Naechster")
    second = find_paragraph(document, "Zweiter Absatz")
    assert numbering_of(second) is None, "Folgeabsatz bekommt kein Aufzaehlungszeichen"
    assert second.paragraph_format.left_indent is not None


def test_code_block_inside_list(doc, assert_valid, convert):
    path, _ = convert("- Punkt mit Code\n\n  ```python\n  x = 1\n  ```\n\n- Weiter")
    assert_valid(path)
    from docx import Document

    document = Document(str(path))
    code = find_paragraph(document, "x = 1")
    assert code.style.style_id == st.S_CODE_BLOCK


# ----------------------------------------------------------------------
# Zitate
# ----------------------------------------------------------------------
def test_blockquote_style_and_indent(doc):
    document = doc("> Ein Zitat")
    paragraph = find_paragraph(document, "Ein Zitat")
    assert paragraph.style.style_id == st.S_QUOTE
    # Word speichert in Twips, daher Rundungstoleranz
    assert abs(paragraph.paragraph_format.left_indent - Mm(8.0)) < 700


def test_nested_blockquote_indents_further(doc):
    document = doc("> Aussen\n>\n> > Innen")
    outer = find_paragraph(document, "Aussen").paragraph_format.left_indent
    inner = find_paragraph(document, "Innen").paragraph_format.left_indent
    assert inner > outer


def test_blockquote_keeps_heading_style(doc):
    document = doc("> ## Zitierte Ueberschrift")
    paragraph = find_paragraph(document, "Zitierte Ueberschrift")
    assert paragraph.style.style_id == "Heading2"


# ----------------------------------------------------------------------
# Code
# ----------------------------------------------------------------------
def test_code_block_one_paragraph_per_line(doc):
    document = doc("```\neins\nzwei\ndrei\n```")
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    assert [p.text for p in code] == ["eins", "zwei", "drei"]


def test_code_block_highlighting_colors_runs(doc):
    document = doc("```python\ndef f():\n    return 1\n```")
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    colored = [r for p in code for r in p.runs if r.font.color and r.font.color.rgb]
    assert colored, "mit Pygments muessen Schluesselwoerter eingefaerbt sein"


def test_highlighting_can_be_disabled(doc):
    document = doc("```python\ndef f():\n    pass\n```", highlight=False)
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    colored = [r for p in code for r in p.runs if r.font.color and r.font.color.rgb]
    assert not colored


def test_code_block_preserves_indentation(doc):
    document = doc("```\n    eingerueckt\n```")
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    assert code[0].text == "    eingerueckt"


def test_code_block_outer_borders_only(doc):
    document = doc("```\na\nb\nc\n```")
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    first = code[0]._p.find(w("pPr")).find(w("pBdr"))
    middle = code[1]._p.find(w("pPr")).find(w("pBdr"))
    last = code[2]._p.find(w("pPr")).find(w("pBdr"))
    assert first.find(w("top")) is not None
    assert middle.find(w("top")) is None and middle.find(w("bottom")) is None
    assert last.find(w("bottom")) is not None


def test_unknown_language_falls_back(doc):
    document = doc("```gibtsnicht\nirgendwas\n```")
    code = [p for p in document.paragraphs if p.style.style_id == st.S_CODE_BLOCK]
    assert [p.text for p in code] == ["irgendwas"]


# ----------------------------------------------------------------------
# Tabellen
# ----------------------------------------------------------------------
def test_table_dimensions_and_header(doc):
    document = doc("| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |")
    table = document.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 2)
    assert table.rows[0].cells[0].text == "A"
    assert table.rows[2].cells[1].text == "4"


def test_table_header_is_bold_and_shaded(doc):
    document = doc("| Kopf |\n|---|\n| Zelle |")
    header = document.tables[0].rows[0].cells[0]
    assert all(r.bold for p in header.paragraphs for r in p.runs)
    shd = header._tc.find(w("tcPr")).find(w("shd"))
    assert shd is not None and shd.get(w("fill")) != "auto"


def test_table_header_repeats_on_new_page(doc):
    document = doc("| Kopf |\n|---|\n| Zelle |")
    trPr = document.tables[0].rows[0]._tr.find(w("trPr"))
    assert trPr.find(w("tblHeader")) is not None


def test_table_column_alignment(doc):
    document = doc("| L | R |\n|:--|--:|\n| a | b |")
    row = document.tables[0].rows[1]
    left = row.cells[0].paragraphs[0].alignment
    right = row.cells[1].paragraphs[0].alignment
    assert str(left) != str(right)


def test_table_cell_inline_formatting(doc):
    document = doc("| A |\n|---|\n| **fett** |")
    cell = document.tables[0].rows[1].cells[0]
    assert any(r.bold for p in cell.paragraphs for r in p.runs)


def test_ragged_table_does_not_crash(doc, assert_valid, convert):
    path, _ = convert("| A | B | C |\n|---|---|---|\n| 1 |\n| 1 | 2 | 3 | 4 |")
    assert_valid(path)


def test_table_widths_fit_page(doc):
    document = doc("| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |")
    table = document.tables[0]
    total = sum(cell.width.mm for cell in table.rows[0].cells)
    assert 155 < total < 165, f"Summe der Spaltenbreiten: {total} mm (Textbreite 160 mm)"


# ----------------------------------------------------------------------
# Sonstige Bloecke
# ----------------------------------------------------------------------
def test_horizontal_rule(doc):
    document = doc("oben\n\n---\n\nunten")
    rules = [p for p in document.paragraphs if p.style.style_id == st.S_HRULE]
    assert len(rules) == 1


def test_definition_list(doc):
    document = doc("Begriff\n: Die Erklaerung")
    assert find_paragraph(document, "Begriff").style.style_id == st.S_DEF_TERM
    assert find_paragraph(document, "Die Erklaerung").style.style_id == st.S_DEF_BODY


@pytest.mark.parametrize(
    "marker", ["<!-- pagebreak -->", "\\newpage", "{{pagebreak}}", "<!-- page-break -->"]
)
def test_page_break_markers(doc, marker):
    document = doc(f"oben\n\n{marker}\n\nunten")
    breaks = [n for n in document.element.body.iter(w("br")) if n.get(w("type")) == "page"]
    assert len(breaks) == 1, f"Marker {marker!r} nicht erkannt"
