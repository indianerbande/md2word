"""Uebertraegt den HTML-Baum des Markdown-Dokuments in ein Word-Dokument."""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from typing import Any, Iterable

import lxml.html
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from md2word import docxutil as dx
from md2word import images as img
from md2word import styles as st
from md2word.config import Config
from md2word.highlight import highlight_code

_INLINE_TAGS = {
    "a", "b", "strong", "i", "em", "u", "s", "del", "strike", "code", "span",
    "sup", "sub", "br", "img", "mark", "kbd", "samp", "var", "abbr", "cite",
    "small", "big", "tt", "q", "ins", "label", "input", "wbr", "time",
}

_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_CHECKBOX_CHECKED = "☒ "   # gekreuztes Kaestchen
_CHECKBOX_OPEN = "☐ "      # leeres Kaestchen


@dataclass
class InlineFormat:
    """Zeichenformatierung, die beim Abstieg durch Inline-Tags mitwaechst."""

    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    code: bool = False
    superscript: bool = False
    subscript: bool = False
    small_caps: bool = False
    color: str | None = None
    highlight: str | None = None
    font: str | None = None
    size: float | None = None
    link: str | None = None

    def merged(self, **changes: Any) -> "InlineFormat":
        data = self.__dict__.copy()
        data.update(changes)
        return InlineFormat(**data)


@dataclass
class RenderState:
    """Wo im Dokument gerade geschrieben wird."""

    container: Any                      # Document, _Cell oder Footnote-Kontext
    list_stack: list[tuple[int, bool]] = field(default_factory=list)
    quote_depth: int = 0
    indent_mm: float = 0.0
    in_footnote: bool = False


class Warnings:
    """Sammelt nicht-fatale Probleme, damit sie gebuendelt gemeldet werden."""

    def __init__(self, verbose: bool = False) -> None:
        self.items: list[str] = []
        self.verbose = verbose

    def add(self, message: str) -> None:
        if message not in self.items:
            self.items.append(message)
        if self.verbose:
            print(f"  ! {message}", file=sys.stderr)


class DocxRenderer:
    """Baut aus dem HTML-Baum das Word-Dokument auf."""

    def __init__(self, document: Any, config: Config) -> None:
        self.doc = document
        self.cfg = config
        self.styles = st.StyleLookup(document)
        self.numbering = dx.NumberingRegistry(document)
        self.warnings = Warnings(config.verbose)

        self._bookmark_id = 1000
        self._anchors: set[str] = set()
        self._heading_anchors: dict[str, str] = {}
        self._footnote_html: dict[str, Any] = {}
        self._footnotes: Any = None
        self._endnotes: list[Any] = []
        self._first_block = True

    # ------------------------------------------------------------------
    # Einstieg
    # ------------------------------------------------------------------
    def render(self, html: str) -> None:
        root = lxml.html.fragment_fromstring(html or "", create_parent="div")
        self._collect_footnotes(root)
        self._prescan_headings(root)

        if self.cfg.footnote_mode == "footnotes" and self._footnote_html:
            try:
                self._footnotes = dx.FootnoteStore(self.doc)
            except Exception as exc:  # pragma: no cover - defensiv
                self.warnings.add(f"Echte Fussnoten nicht moeglich ({exc}) - nutze Endnoten")
                self.cfg.footnote_mode = "endnotes"

        state = RenderState(container=self.doc)
        self._render_children(root, state)
        self._emit_endnotes()

    # ------------------------------------------------------------------
    # Vorbereitung
    # ------------------------------------------------------------------
    def _collect_footnotes(self, root: Any) -> None:
        """Schneidet den Fussnoten-Abschnitt heraus und merkt sich die Inhalte."""
        sections = root.xpath(".//section[contains(@class,'footnotes')]")
        for section in sections:
            for item in section.xpath(".//li"):
                item_id = item.get("id")
                if item_id:
                    for backref in item.xpath(".//a[contains(@class,'footnote-backref')]"):
                        _drop(backref)
                    self._footnote_html[item_id] = item
            _drop(section)

        for separator in root.xpath(".//hr[contains(@class,'footnotes-sep')]"):
            _drop(separator)

    def _prescan_headings(self, root: Any) -> None:
        """Vergibt Anker-IDs fuer Ueberschriften, damit [Text](#anker) klappt."""
        for level in range(1, 7):
            for heading in root.xpath(f".//h{level}"):
                text = "".join(heading.itertext()).strip()
                explicit = heading.get("id")
                anchor = explicit or dx.slugify(text, self._anchors)
                if explicit:
                    self._anchors.add(explicit)
                heading.set("data-md2word-anchor", anchor)
                self._heading_anchors[anchor] = text

    # ------------------------------------------------------------------
    # Blockebene
    # ------------------------------------------------------------------
    def _render_children(self, element: Any, state: RenderState) -> None:
        if element.text and element.text.strip():
            self._loose_text(element.text, state)
        for child in element:
            self._render_block(child, state)
            if child.tail and child.tail.strip():
                self._loose_text(child.tail, state)

    def _loose_text(self, text: str, state: RenderState) -> None:
        """Text direkt auf Blockebene (kommt bei rohem HTML vor)."""
        cleaned = _collapse(text)
        if not cleaned:
            return
        paragraph = self._new_paragraph(state)
        self._add_run(paragraph, cleaned, InlineFormat(), state)

    def _render_block(self, node: Any, state: RenderState) -> None:
        tag = _tag(node)

        if tag is None:  # Kommentar oder Processing Instruction
            return

        if tag in _INLINE_TAGS:
            paragraph = self._new_paragraph(state)
            self._render_inline(node, paragraph, InlineFormat(), state, include_tail=False)
            self._drop_if_empty(paragraph)
            return

        handler = {
            "p": self._block_paragraph,
            "h1": self._block_heading, "h2": self._block_heading,
            "h3": self._block_heading, "h4": self._block_heading,
            "h5": self._block_heading, "h6": self._block_heading,
            "ul": self._block_list, "ol": self._block_list,
            "blockquote": self._block_quote,
            "pre": self._block_pre,
            "table": self._block_table,
            "hr": self._block_hr,
            "dl": self._block_deflist,
            "figure": self._block_container,
            "figcaption": self._block_figcaption,
            "section": self._block_container,
            "article": self._block_container,
            "main": self._block_container,
            "header": self._block_container,
            "footer": self._block_container,
            "aside": self._block_container,
            "details": self._block_container,
            "summary": self._block_paragraph,
            "div": self._block_div,
            "li": self._block_container,
            "script": _ignore,
            "style": _ignore,
            "head": _ignore,
            "meta": _ignore,
            "link": _ignore,
            "title": _ignore,
        }.get(tag)

        if handler is None:
            if self.cfg.strip_html:
                return
            self._block_container(node, state)
            return

        handler(node, state)

    def _block_container(self, node: Any, state: RenderState) -> None:
        """Unbekannter Block: Inhalt uebernehmen, Huelle ignorieren."""
        if _has_block_children(node):
            self._render_children(node, state)
        else:
            self._block_paragraph(node, state)

    def _block_div(self, node: Any, state: RenderState) -> None:
        classes = (node.get("class") or "").split()
        if "md2word-pagebreak" in classes:
            paragraph = self._new_paragraph(state)
            paragraph.style = self.styles["Normal"]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            return
        if "math" in classes:
            self._block_math(node, state)
            return
        self._block_container(node, state)

    def _block_math(self, node: Any, state: RenderState) -> None:
        formula = _collapse("".join(node.itertext()))
        if not formula:
            return
        paragraph = self._new_paragraph(state)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(formula)
        run.italic = True
        dx.set_run_font(run, "Cambria Math")
        self.warnings.add(
            "Mathematischer Ausdruck wurde als formatierter Text uebernommen, "
            "nicht als Word-Formel"
        )

    def _block_paragraph(self, node: Any, state: RenderState) -> None:
        paragraph = self._new_paragraph(state)
        self._apply_alignment(paragraph, node)
        self._render_inline(node, paragraph, InlineFormat(), state, include_tail=False)
        self._drop_if_empty(paragraph)

    def _block_heading(self, node: Any, state: RenderState) -> None:
        level = int(_tag(node)[1])
        if self.cfg.break_on_h1 and level == 1 and not self._first_block:
            breaker = self._new_paragraph(state)
            breaker.paragraph_format.space_after = Pt(0)
            breaker.add_run().add_break(WD_BREAK.PAGE)

        paragraph = self._new_paragraph(state)
        paragraph.style = self.styles.heading(level)
        self._render_inline(node, paragraph, InlineFormat(), state, include_tail=False)

        anchor = node.get("data-md2word-anchor")
        if anchor:
            self._bookmark_id += 1
            dx.add_bookmark(paragraph, anchor, self._bookmark_id)
        dx.keep_with_next(paragraph)
        self._first_block = False

    def _block_quote(self, node: Any, state: RenderState) -> None:
        inner = RenderState(
            container=state.container,
            list_stack=list(state.list_stack),
            quote_depth=state.quote_depth + 1,
            indent_mm=state.indent_mm,
        )
        start = self._paragraph_count(state)
        self._render_children(node, inner)

        for paragraph in self._paragraphs_since(state, start):
            style_id = paragraph.style.style_id
            if style_id == st.S_CODE_BLOCK or style_id.startswith("Heading"):
                continue
            # Absaetze eines verschachtelten Zitats sind bereits fertig
            # eingerueckt - sie hier erneut anzufassen wuerde die Staffelung
            # wieder einebnen.
            if style_id == st.S_QUOTE:
                continue
            paragraph.style = self.styles[st.S_QUOTE]
            paragraph.paragraph_format.left_indent = Mm(8.0 * inner.quote_depth)

    def _block_pre(self, node: Any, state: RenderState) -> None:
        code_node = node.find("code")
        source_node = code_node if code_node is not None else node
        code = "".join(source_node.itertext())
        code = code.rstrip("\n")

        language = ""
        classes = (source_node.get("class") or "").split()
        for cls in classes:
            if cls.startswith("language-"):
                language = cls[len("language-"):]
            elif cls.startswith("lang-"):
                language = cls[len("lang-"):]

        self._emit_code_block(code, language, state)

    def _emit_code_block(self, code: str, language: str, state: RenderState) -> None:
        lines = code.split("\n") or [""]
        fragments = (
            highlight_code(code, language, self.cfg.pygments_style)
            if self.cfg.highlight
            else None
        )
        per_line = _fragments_by_line(fragments, len(lines)) if fragments else None

        paragraphs: list[Paragraph] = []
        for index, line in enumerate(lines):
            paragraph = self._new_paragraph(state)
            paragraph.style = self.styles[st.S_CODE_BLOCK]
            if state.indent_mm:
                paragraph.paragraph_format.left_indent = Mm(3.0 + state.indent_mm)

            pieces = per_line[index] if per_line else [(line, None, False, False)]
            if not pieces and line == "":
                paragraph.add_run("")
            for text, color, bold, italic in pieces:
                if not text:
                    continue
                run = paragraph.add_run(text.replace("\t", "    "))
                if color:
                    run.font.color.rgb = st.hex_to_rgb(color)
                run.bold = bold or None
                run.italic = italic or None
            paragraphs.append(paragraph)

        if not paragraphs:
            return

        # Rahmen nur aussen: oben am ersten, unten am letzten Absatz
        for position, paragraph in enumerate(paragraphs):
            edges = ["left", "right"]
            if position == 0:
                edges.append("top")
            if position == len(paragraphs) - 1:
                edges.append("bottom")
            dx.set_paragraph_border(
                paragraph, edges=tuple(edges), size=4, space=6, color="D0D0D0"
            )
        paragraphs[0].paragraph_format.space_before = Pt(6)
        paragraphs[-1].paragraph_format.space_after = Pt(10)

    def _block_list(self, node: Any, state: RenderState) -> None:
        ordered = _tag(node) == "ol"
        start_at = _int_or(node.get("start"), 1)
        level = len(state.list_stack)
        num_id = self.numbering.new_list(ordered, start_at)

        child_state = RenderState(
            container=state.container,
            list_stack=state.list_stack + [(num_id, ordered)],
            quote_depth=state.quote_depth,
            indent_mm=state.indent_mm,
        )

        for item in node.xpath("./li"):
            self._render_list_item(item, child_state, num_id, level)

    def _render_list_item(self, item: Any, state: RenderState, num_id: int, level: int) -> None:
        checkbox = item.xpath(".//input[@type='checkbox']")
        prefix = ""
        if checkbox:
            prefix = _CHECKBOX_CHECKED if checkbox[0].get("checked") is not None else _CHECKBOX_OPEN
            for box in checkbox:
                _drop(box, keep_tail=True)
            for label in item.xpath(".//label"):
                _unwrap(label)

        blocks = [child for child in item if _tag(child) in {"ul", "ol"}]
        content_nodes = [child for child in item if child not in blocks]

        first = self._new_paragraph(state)
        first.style = self.styles["ListParagraph"]
        dx.apply_numbering(first, num_id, level)
        first.paragraph_format.space_after = Pt(2)
        if state.indent_mm:
            base = first.paragraph_format.left_indent or Mm(0)
            first.paragraph_format.left_indent = base + Mm(state.indent_mm)

        if prefix:
            first.add_run(prefix)

        # Direkter Text und Inline-Inhalt landen im ersten Absatz
        if item.text and item.text.strip():
            self._add_run(first, _collapse_leading(item.text), InlineFormat(), state)

        # Nur der erste Absatz des Listenpunkts traegt das Aufzaehlungszeichen;
        # alles Weitere wird darunter eingerueckt.
        pending_blocks: list[Any] = []
        first_paragraph_used = False
        for child in content_nodes:
            tag = _tag(child)
            if pending_blocks:
                pending_blocks.append(child)
            elif tag == "p" and not first_paragraph_used:
                self._render_inline(child, first, InlineFormat(), state, include_tail=False)
                if child.tail and child.tail.strip():
                    self._add_run(first, _collapse(child.tail), InlineFormat(), state)
                first_paragraph_used = True
            elif tag in _INLINE_TAGS:
                self._render_inline_node(child, first, InlineFormat(), state)
            else:
                pending_blocks.append(child)

        self._drop_if_empty(first, force_keep=bool(prefix))

        # Weitere Bloecke im selben Listenpunkt: eingerueckt, ohne Aufzaehlungszeichen
        if pending_blocks:
            nested_state = RenderState(
                container=state.container,
                list_stack=state.list_stack,
                quote_depth=state.quote_depth,
                indent_mm=state.indent_mm + 7.4 * (level + 1),
            )
            for child in pending_blocks:
                start = self._paragraph_count(state)
                self._render_block(child, nested_state)
                for paragraph in self._paragraphs_since(state, start):
                    if paragraph.style.style_id == st.S_CODE_BLOCK:
                        continue
                    current = paragraph.paragraph_format.left_indent or Mm(0)
                    if current < Mm(nested_state.indent_mm):
                        paragraph.paragraph_format.left_indent = Mm(nested_state.indent_mm)

        # Verschachtelte Listen
        for sublist in blocks:
            self._block_list(sublist, state)

    def _block_deflist(self, node: Any, state: RenderState) -> None:
        for child in node:
            tag = _tag(child)
            if tag == "dt":
                paragraph = self._new_paragraph(state)
                paragraph.style = self.styles[st.S_DEF_TERM]
                self._render_inline(child, paragraph, InlineFormat(), state, include_tail=False)
            elif tag == "dd":
                if _has_block_children(child):
                    start = self._paragraph_count(state)
                    self._render_children(child, state)
                    for paragraph in self._paragraphs_since(state, start):
                        paragraph.paragraph_format.left_indent = Mm(8.0)
                else:
                    paragraph = self._new_paragraph(state)
                    paragraph.style = self.styles[st.S_DEF_BODY]
                    self._render_inline(child, paragraph, InlineFormat(), state, include_tail=False)

    def _block_hr(self, node: Any, state: RenderState) -> None:
        paragraph = self._new_paragraph(state)
        paragraph.style = self.styles[st.S_HRULE]

    def _block_figcaption(self, node: Any, state: RenderState) -> None:
        paragraph = self._new_paragraph(state)
        paragraph.style = self.styles[st.S_CAPTION]
        self._render_inline(node, paragraph, InlineFormat(), state, include_tail=False)

    # ------------------------------------------------------------------
    # Tabellen
    # ------------------------------------------------------------------
    def _block_table(self, node: Any, state: RenderState) -> None:
        rows = node.xpath("./thead/tr") + node.xpath("./tbody/tr") + node.xpath("./tr")
        if not rows:
            return

        head_count = len(node.xpath("./thead/tr"))
        columns = max(len(row.xpath("./th|./td")) for row in rows)
        if columns == 0:
            return

        target = state.container if isinstance(state.container, _Cell) else self.doc
        table = target.add_table(rows=len(rows), cols=columns)
        table.style = self.styles[st.build_table_style(self.doc, self.cfg)]
        st.apply_table_alignment(table)
        dx.set_cell_margins(table)

        widths = self._column_widths(rows, columns, state)
        dx.set_table_layout_fixed(table)

        for row_index, row_node in enumerate(rows):
            cells = row_node.xpath("./th|./td")
            is_header = row_index < head_count or all(_tag(c) == "th" for c in cells)
            docx_row = table.rows[row_index]
            dx.prevent_row_split(docx_row)
            if is_header:
                dx.set_repeat_header(docx_row)

            for col_index in range(columns):
                cell = docx_row.cells[col_index]
                cell.width = widths[col_index]
                dx.set_cell_vertical_alignment(cell, "center")
                if is_header:
                    dx.set_cell_background(cell, self.cfg.table_header_bg)

                if col_index >= len(cells):
                    _clear_cell(cell)
                    continue

                self._fill_cell(cell, cells[col_index], state, bold=is_header)

        # Leerabsatz nach der Tabelle: sonst kleben zwei Tabellen aneinander
        if not isinstance(state.container, _Cell):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

    def _column_widths(self, rows: Iterable[Any], columns: int, state: RenderState) -> list[Mm]:
        """Verteilt die Textbreite proportional zur Inhaltslaenge auf die Spalten."""
        weights = [1.0] * columns
        for row in rows:
            cells = row.xpath("./th|./td")
            for index, cell in enumerate(cells[:columns]):
                length = len(_collapse("".join(cell.itertext())))
                weights[index] = max(weights[index], min(float(length), 60.0))

        total = sum(weights) or float(columns)
        available = self.cfg.text_width_mm - state.indent_mm
        if isinstance(state.container, _Cell):
            available = max(30.0, available * 0.9)

        average = available / columns
        result = []
        for weight in weights:
            share = available * weight / total
            share = max(average * 0.45, min(share, average * 2.2))
            result.append(share)

        scale = available / sum(result)
        return [Mm(round(value * scale, 2)) for value in result]

    def _fill_cell(self, cell: _Cell, node: Any, state: RenderState, bold: bool) -> None:
        _clear_cell(cell)
        cell_state = RenderState(
            container=cell,
            list_stack=[],
            quote_depth=0,
            indent_mm=0.0,
        )

        align = _style_alignment(node)
        if _has_block_children(node):
            self._render_children(node, cell_state)
        else:
            paragraph = cell.add_paragraph()
            paragraph.style = self.styles[st.S_TABLE_TEXT]
            self._render_inline(
                node, paragraph, InlineFormat(bold=bold), cell_state, include_tail=False
            )

        for paragraph in cell.paragraphs:
            if paragraph.style.style_id == "Normal":
                paragraph.style = self.styles[st.S_TABLE_TEXT]
            if align:
                paragraph.alignment = align
            if bold:
                for run in paragraph.runs:
                    run.bold = True

        if not cell.paragraphs:
            cell.add_paragraph()

    # ------------------------------------------------------------------
    # Inline-Ebene
    # ------------------------------------------------------------------
    def _render_inline(
        self,
        node: Any,
        paragraph: Paragraph,
        fmt: InlineFormat,
        state: RenderState,
        include_tail: bool = True,
    ) -> None:
        if node.text:
            self._add_run(paragraph, node.text, fmt, state)
        for child in node:
            self._render_inline_node(child, paragraph, fmt, state)
        if include_tail and node.tail:
            self._add_run(paragraph, node.tail, fmt, state)

    def _render_inline_node(
        self, node: Any, paragraph: Paragraph, fmt: InlineFormat, state: RenderState
    ) -> None:
        tag = _tag(node)
        if tag is None:
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        classes = (node.get("class") or "").split()

        if tag == "br":
            paragraph.add_run().add_break()
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        if tag == "img":
            self._inline_image(node, paragraph, state)
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        if tag == "input":
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        if tag == "sup" and "footnote-ref" in classes:
            self._inline_footnote(node, paragraph, state)
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        if tag == "span" and "math" in classes:
            text = _collapse("".join(node.itertext()))
            if text:
                run = paragraph.add_run(text)
                run.italic = True
                dx.set_run_font(run, "Cambria Math")
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        if tag == "a":
            self._inline_link(node, paragraph, fmt, state)
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        child_fmt = self._extend_format(fmt, tag, node)

        if tag in {"ul", "ol", "p", "div", "blockquote", "pre", "table"}:
            # Blockelement mitten im Inline-Kontext: als Text uebernehmen
            text = _collapse("".join(node.itertext()))
            if text:
                self._add_run(paragraph, text, child_fmt, state)
            if node.tail:
                self._add_run(paragraph, node.tail, fmt, state)
            return

        self._render_inline(node, paragraph, child_fmt, state, include_tail=False)
        if node.tail:
            self._add_run(paragraph, node.tail, fmt, state)

    def _extend_format(self, fmt: InlineFormat, tag: str, node: Any) -> InlineFormat:
        changes: dict[str, Any] = {}
        if tag in {"strong", "b"}:
            changes["bold"] = True
        elif tag in {"em", "i", "cite", "var"}:
            changes["italic"] = True
        elif tag in {"u", "ins"}:
            changes["underline"] = True
        elif tag in {"s", "del", "strike"}:
            changes["strike"] = True
        elif tag in {"code", "kbd", "samp", "tt"}:
            changes["code"] = True
        elif tag == "sup":
            changes["superscript"] = True
        elif tag == "sub":
            changes["subscript"] = True
        elif tag == "mark":
            changes["highlight"] = "FFFF00"
        elif tag == "small":
            changes["size"] = max(6.0, self.cfg.font_size - 2)
        elif tag == "abbr":
            changes["small_caps"] = True

        style = node.get("style") or ""
        color = re.search(r"(?<!-)\bcolor\s*:\s*#?([0-9A-Fa-f]{6})\b", style)
        if color:
            changes["color"] = color.group(1).upper()
        if "font-weight:bold" in style.replace(" ", ""):
            changes["bold"] = True
        if "font-style:italic" in style.replace(" ", ""):
            changes["italic"] = True

        return fmt.merged(**changes) if changes else fmt

    def _add_run(
        self, paragraph: Paragraph, text: str, fmt: InlineFormat, state: RenderState
    ) -> Any:
        if text is None:
            return None
        content = text if fmt.code else _collapse_soft(text)
        if not content:
            return None

        run = paragraph.add_run(content)
        self._apply_format(run, fmt)
        return run

    def _apply_format(self, run: Any, fmt: InlineFormat) -> None:
        if fmt.bold:
            run.bold = True
        if fmt.italic:
            run.italic = True
        if fmt.underline:
            run.underline = True
        if fmt.strike:
            run.font.strike = True
        if fmt.small_caps:
            run.font.small_caps = True
        if fmt.superscript:
            dx.set_vertical_align(run, "superscript")
        if fmt.subscript:
            dx.set_vertical_align(run, "subscript")
        if fmt.code:
            dx.set_run_style(run, st.S_CODE_INLINE)
            dx.set_run_font(run, self.cfg.code_font)
            run.font.size = Pt(self.cfg.code_font_size)
            dx.set_no_proof(run)
        if fmt.color:
            run.font.color.rgb = st.hex_to_rgb(fmt.color)
        if fmt.highlight:
            dx.set_shading(run._r.get_or_add_rPr(), fmt.highlight)
        if fmt.font:
            dx.set_run_font(run, fmt.font)
        if fmt.size:
            run.font.size = Pt(fmt.size)
        if fmt.link:
            dx.set_run_style(run, st.S_HYPERLINK)
            run.font.color.rgb = st.hex_to_rgb(self.cfg.link_color)
            run.underline = True

    def _inline_link(
        self, node: Any, paragraph: Paragraph, fmt: InlineFormat, state: RenderState
    ) -> None:
        href = (node.get("href") or "").strip()
        title = node.get("title") or ""
        link_fmt = fmt.merged(link=href or None)

        if not href:
            self._render_inline(node, paragraph, fmt, state, include_tail=False)
            return

        # Interner Anker
        if href.startswith("#"):
            anchor = href[1:]
            if anchor not in self._anchors:
                self.warnings.add(f"Interner Verweis ohne Ziel: {href}")
                self._render_inline(node, paragraph, fmt, state, include_tail=False)
                return
            container = dx.add_internal_hyperlink(paragraph, anchor)
        else:
            try:
                container = dx.add_external_hyperlink(paragraph, href)
            except Exception as exc:
                self.warnings.add(f"Link konnte nicht gesetzt werden ({href}): {exc}")
                self._render_inline(node, paragraph, link_fmt, state, include_tail=False)
                return

        before = len(paragraph.runs)
        self._render_inline(node, paragraph, link_fmt, state, include_tail=False)
        created = paragraph.runs[before:]
        if not created:
            created = [self._add_run(paragraph, href, link_fmt, state)]
        for run in created:
            if run is not None:
                dx.move_run_into(container, run)

        if title and title != _collapse("".join(node.itertext())):
            self._add_run(paragraph, f" ({title})", fmt, state)

    def _inline_footnote(self, node: Any, paragraph: Paragraph, state: RenderState) -> None:
        anchor = node.xpath(".//a/@href")
        key = anchor[0].lstrip("#") if anchor else ""
        source = self._footnote_html.get(key)

        if source is None:
            text = _collapse("".join(node.itertext()))
            run = paragraph.add_run(text)
            dx.set_vertical_align(run, "superscript")
            return

        if self.cfg.footnote_mode == "footnotes" and self._footnotes is not None:
            self._emit_real_footnote(source, paragraph)
        else:
            self._emit_endnote_reference(source, paragraph)

    def _emit_real_footnote(self, source: Any, paragraph: Paragraph) -> None:
        footnote_id, first_p = self._footnotes.add_footnote()
        self._footnotes.add_reference(paragraph, footnote_id)

        blocks = [child for child in source if _tag(child) not in (None,)]
        note_state = RenderState(container=None, in_footnote=True)

        target = Paragraph(first_p, self._footnotes._part)
        wrote_any = False

        if source.text and source.text.strip():
            self._add_run(target, source.text, InlineFormat(), note_state)
            wrote_any = True

        for index, block in enumerate(blocks):
            if index > 0:
                extra = self._footnotes.add_paragraph_to(footnote_id)
                target = Paragraph(extra, self._footnotes._part)
            if _tag(block) in _INLINE_TAGS:
                self._render_inline_node(block, target, InlineFormat(), note_state)
            else:
                self._render_inline(block, target, InlineFormat(), note_state, include_tail=False)
            wrote_any = True

        if not wrote_any:
            self._add_run(target, _collapse("".join(source.itertext())), InlineFormat(), note_state)

    def _emit_endnote_reference(self, source: Any, paragraph: Paragraph) -> None:
        self._endnotes.append(source)
        run = paragraph.add_run(f"[{len(self._endnotes)}]")
        dx.set_vertical_align(run, "superscript")

    def _emit_endnotes(self) -> None:
        if not self._endnotes:
            return
        heading = self.doc.add_paragraph()
        heading.style = self.styles["Heading2"]
        heading.add_run("Anmerkungen")

        num_id = self.numbering.new_list(ordered=True)
        state = RenderState(container=self.doc)
        for source in self._endnotes:
            paragraph = self.doc.add_paragraph()
            paragraph.style = self.styles["ListParagraph"]
            dx.apply_numbering(paragraph, num_id, 0)
            paragraph.paragraph_format.space_after = Pt(2)
            blocks = list(source)
            if blocks:
                for block in blocks:
                    self._render_inline(block, paragraph, InlineFormat(), state, include_tail=False)
            else:
                self._add_run(paragraph, "".join(source.itertext()), InlineFormat(), state)

    # ------------------------------------------------------------------
    # Bilder
    # ------------------------------------------------------------------
    def _inline_image(self, node: Any, paragraph: Paragraph, state: RenderState) -> None:
        src = node.get("src") or ""
        alt = (node.get("alt") or "").strip()
        title = (node.get("title") or "").strip()

        try:
            loaded = img.load_image(
                src,
                base_dir=self.cfg.base_dir,
                download=self.cfg.download_images,
                timeout=self.cfg.image_timeout,
            )
        except Exception as exc:
            self.warnings.add(f"Bild uebersprungen: {exc}")
            fallback = alt or title or src
            if fallback:
                run = paragraph.add_run(f"[Bild: {fallback}]")
                run.italic = True
                run.font.color.rgb = st.hex_to_rgb(self.cfg.quote_color)
            return

        max_width = self.cfg.max_image_width or (self.cfg.text_width_mm - state.indent_mm)
        explicit = _parse_length(node.get("width"))
        width = Mm(min(explicit, max_width)) if explicit else img.fit_width(loaded, max_width)

        try:
            run = paragraph.add_run()
            run.add_picture(loaded.stream, width=width)
        except Exception as exc:
            self.warnings.add(f"Bild konnte nicht eingebettet werden ({src}): {exc}")
            return

        caption = self._caption_for(alt, title)
        if caption and _is_standalone_image(node):
            paragraph.style = self.styles[st.S_FIGURE]
            caption_paragraph = self._new_paragraph(state)
            caption_paragraph.style = self.styles[st.S_CAPTION]
            caption_paragraph.add_run(caption)

    def _caption_for(self, alt: str, title: str) -> str:
        """Bildunterschrift gemaess --captions bestimmen."""
        mode = self.cfg.captions
        if mode == "none":
            return ""
        if mode == "alt":
            return title or alt
        return title

    # ------------------------------------------------------------------
    # Hilfsfunktionen fuer Absaetze
    # ------------------------------------------------------------------
    def _new_paragraph(self, state: RenderState) -> Paragraph:
        container = state.container if state.container is not None else self.doc
        paragraph = container.add_paragraph()
        if state.indent_mm:
            paragraph.paragraph_format.left_indent = Mm(state.indent_mm)
        self._first_block = False
        return paragraph

    def _paragraph_count(self, state: RenderState) -> int:
        container = state.container if state.container is not None else self.doc
        return len(container.paragraphs)

    def _paragraphs_since(self, state: RenderState, start: int) -> list[Paragraph]:
        container = state.container if state.container is not None else self.doc
        return list(container.paragraphs[start:])

    def _apply_alignment(self, paragraph: Paragraph, node: Any) -> None:
        align = _style_alignment(node)
        if align is not None:
            paragraph.alignment = align

    @staticmethod
    def _drop_if_empty(paragraph: Paragraph, force_keep: bool = False) -> None:
        if force_keep:
            return
        if paragraph.runs or paragraph._p.findall(qn("w:hyperlink")):
            return
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)


# ----------------------------------------------------------------------
# Modul-Hilfsfunktionen
# ----------------------------------------------------------------------
def _tag(node: Any) -> str | None:
    tag = getattr(node, "tag", None)
    if not isinstance(tag, str):
        return None
    return tag.lower()


def _ignore(node: Any, state: RenderState) -> None:
    return None


def _drop(node: Any, keep_tail: bool = False) -> None:
    parent = node.getparent()
    if parent is None:
        return
    if keep_tail and node.tail:
        previous = node.getprevious()
        if previous is not None:
            previous.tail = (previous.tail or "") + node.tail
        else:
            parent.text = (parent.text or "") + node.tail
    parent.remove(node)


def _unwrap(node: Any) -> None:
    """Entfernt ein Element, behaelt aber dessen Inhalt an gleicher Stelle."""
    parent = node.getparent()
    if parent is None:
        return
    index = parent.index(node)
    previous = node.getprevious()

    if node.text:
        if previous is not None:
            previous.tail = (previous.tail or "") + node.text
        else:
            parent.text = (parent.text or "") + node.text

    for offset, child in enumerate(list(node)):
        parent.insert(index + offset, child)

    if node.tail:
        children = list(node)
        if children:
            last = children[-1]
            last.tail = (last.tail or "") + node.tail
        elif previous is not None:
            previous.tail = (previous.tail or "") + node.tail
        else:
            parent.text = (parent.text or "") + node.tail

    parent.remove(node)


def _has_block_children(node: Any) -> bool:
    return any(_tag(child) not in _INLINE_TAGS and _tag(child) is not None for child in node)


# Geschuetzte Leerzeichen ueberleben das Zusammenfassen: Browser behandeln
# &nbsp; ebenso, und die franzoesische Typografie braucht das schmale
# geschuetzte Leerzeichen vor dem schliessenden Guillemet. Als Escapes
# geschrieben - im Quelltext waeren die Zeichen unsichtbar.
PROTECTED_SPACES = "\u00a0\u202f\u2007"  # NBSP, schmales NBSP, Ziffernbreite
_COLLAPSIBLE = re.compile(rf"[^\S{PROTECTED_SPACES}]+")
_EDGE_SPACE = re.compile(rf"^[^\S{PROTECTED_SPACES}]+|[^\S{PROTECTED_SPACES}]+$")
_LEADING_SPACE = re.compile(rf"^[^\S{PROTECTED_SPACES}]+")


def _collapse(text: str | None) -> str:
    return _EDGE_SPACE.sub("", _COLLAPSIBLE.sub(" ", text or ""))


def _collapse_soft(text: str | None) -> str:
    """Wie HTML: Folgen von Leerraum werden zu einem Leerzeichen, Raender bleiben."""
    return _COLLAPSIBLE.sub(" ", text or "")


def _collapse_leading(text: str | None) -> str:
    return _LEADING_SPACE.sub("", _COLLAPSIBLE.sub(" ", text or ""))


def _style_alignment(node: Any) -> Any:
    style = (node.get("style") or "").replace(" ", "").lower()
    match = re.search(r"text-align:(\w+)", style)
    if match:
        return _ALIGNMENTS.get(match.group(1))
    align = (node.get("align") or "").lower()
    return _ALIGNMENTS.get(align)


def _int_or(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _parse_length(value: Any) -> float | None:
    """Interpretiert width-Angaben als Millimeter (px werden bei 96 dpi gerechnet)."""
    if not value:
        return None
    text = str(value).strip().lower()
    match = re.match(r"^(\d+(?:\.\d+)?)\s*(px|mm|cm|in|pt)?$", text)
    if not match:
        return None
    number = float(match.group(1))
    unit = match.group(2) or "px"
    factors = {"px": 25.4 / 96, "mm": 1.0, "cm": 10.0, "in": 25.4, "pt": 25.4 / 72}
    return number * factors[unit]


def _is_standalone_image(node: Any) -> bool:
    """True, wenn das Bild allein in seinem Absatz steht."""
    parent = node.getparent()
    if parent is None:
        return False
    if (parent.text or "").strip():
        return False
    siblings = [child for child in parent if child is not node]
    if any(_tag(child) != "br" for child in siblings):
        return False
    return not (node.tail or "").strip()


def _clear_cell(cell: _Cell) -> None:
    for paragraph in list(cell.paragraphs):
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)


def _fragments_by_line(
    fragments: list[Any], line_count: int
) -> list[list[tuple[str, str | None, bool, bool]]]:
    """Verteilt Pygments-Fragmente auf Zeilen (Runs duerfen kein \\n enthalten)."""
    lines: list[list[tuple[str, str | None, bool, bool]]] = [[] for _ in range(line_count)]
    index = 0
    for fragment in fragments:
        parts = fragment.text.split("\n")
        for position, part in enumerate(parts):
            if position > 0:
                index += 1
            if index >= line_count:
                index = line_count - 1
            if part:
                lines[index].append((part, fragment.color, fragment.bold, fragment.italic))
    return lines
