"""Assembles the finished Word document: frame, metadata, content."""

from __future__ import annotations

import datetime as _dt
import os
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

from md2word import docxutil as dx
from md2word import i18n
from md2word import styles as st
from md2word.config import Config
from md2word.parser import markdown_to_html
from md2word.renderer import DocxRenderer

# Front-matter keys that map straight onto configuration fields
_META_KEYS = {
    "title": "title",
    "subtitle": "subtitle",
    "author": "author",
    "authors": "author",
    "date": "date",
    "subject": "subject",
    "description": "subject",
    "keywords": "keywords",
    "tags": "keywords",
    "comments": "comments",
    "lang": "lang",
    "language": "lang",
}

# Front-matter keys that are allowed to switch layout options
_OPTION_KEYS = {
    "toc": bool,
    "toc_depth": int,
    "toc_title": str,
    "title_page": bool,
    "titlepage": bool,
    "number_headings": bool,
    "numbered_headings": bool,
    "page_numbers": bool,
    "header_text": str,
    "footer_text": str,
    "break_on_h1": bool,
    "page_size": str,
    "landscape": bool,
    "theme": str,
    "font_size": float,
    "body_font": str,
    "heading_font": str,
    "code_font": str,
    "highlight": bool,
    "pygments_style": str,
}

_ALIAS = {
    "titlepage": "title_page",
    "numbered_headings": "number_headings",
}


@dataclass
class Result:
    """Result of one conversion."""

    output_path: str
    warnings: list[str]
    heading_count: int


def convert_file(input_path: str, output_path: str, config: Config) -> Result:
    """Reads a Markdown file and writes the Word document."""
    if input_path == "-":
        import sys

        text = sys.stdin.read()
        base_dir = os.getcwd()
    else:
        with open(input_path, "r", encoding="utf-8-sig") as handle:
            text = handle.read()
        base_dir = os.path.dirname(os.path.abspath(input_path)) or "."

    config = _clone_with_base_dir(config, base_dir)
    return convert_text(text, output_path, config, source_name=input_path)


def convert_text(
    text: str, output_path: str, config: Config, source_name: str = ""
) -> Result:
    """Converts Markdown text and saves the result to output_path."""
    html, front_matter = markdown_to_html(text, config)
    config = _apply_front_matter(config, front_matter)

    document = _create_document(config)
    st.apply_styles(document, config, respect_existing=bool(config.reference_doc))

    if config.number_headings:
        dx.enable_heading_numbering(document, config.toc_depth)

    _apply_core_properties(document, config, source_name)

    if config.title_page:
        _build_title_page(document, config)
    elif config.title and not _starts_with_heading(html):
        _build_inline_title(document, config)

    if config.toc:
        _build_toc(document, config)

    renderer = DocxRenderer(document, config)
    renderer.render(html)

    if config.page_numbers or config.header_text or config.footer_text:
        _build_header_footer(document, config)

    if config.toc or config.number_headings:
        dx.force_field_update_on_open(document)

    _remove_leading_empty_paragraph(document)

    dx.trim_paragraph_edges(document.element.body, keep_styles=(st.S_CODE_BLOCK,))
    for part in _footnote_parts(document):
        dx.trim_paragraph_edges(part.element, keep_styles=(st.S_CODE_BLOCK,))

    directory = os.path.dirname(os.path.abspath(output_path))
    if directory:
        os.makedirs(directory, exist_ok=True)
    document.save(output_path)

    heading_count = sum(
        1 for p in document.paragraphs if p.style.style_id.startswith("Heading")
    )
    return Result(
        output_path=output_path,
        warnings=renderer.warnings.items,
        heading_count=heading_count,
    )


# ----------------------------------------------------------------------
def _footnote_parts(document: Any) -> list[Any]:
    """Returns the footnotes part, if one was created."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    return [
        rel.target_part
        for rel in document.part.rels.values()
        if rel.reltype == RT.FOOTNOTES and not rel.is_external
    ]


def _clone_with_base_dir(config: Config, base_dir: str) -> Config:
    from dataclasses import replace

    return replace(config, base_dir=base_dir)


def _create_document(config: Config) -> Any:
    """A new document - optionally based on a reference template."""
    if not config.reference_doc:
        return Document()

    path = os.path.expanduser(config.reference_doc)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"reference document not found: {path}")

    document = Document(path)
    body = document.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)
    return document


def _apply_front_matter(config: Config, meta: dict[str, Any]) -> Config:
    """Applies metadata and layout options from the YAML front matter."""
    if not meta:
        return config

    from dataclasses import replace

    updates: dict[str, Any] = {}
    extra: dict[str, Any] = {}

    for raw_key, value in meta.items():
        key = str(raw_key).strip().lower().replace("-", "_")

        if key in _META_KEYS:
            field = _META_KEYS[key]
            updates[field] = _stringify(value)
            continue

        if key in _OPTION_KEYS:
            field = _ALIAS.get(key, key)
            try:
                updates[field] = _coerce(value, _OPTION_KEYS[key])
            except (TypeError, ValueError):
                extra[key] = value
            continue

        extra[key] = value

    # The command line beats the front matter: explicit values stay
    for field in list(updates):
        if field in config._explicit:
            updates.pop(field)

    updates["_extra"] = {**config._extra, **extra}
    result = replace(config, **updates)
    if "theme" in updates:
        result.apply_theme(result.theme)
    return result


def _stringify(value: Any) -> str:
    if isinstance(value, (list, tuple)):
        return ", ".join(str(item) for item in value)
    if isinstance(value, _dt.date):
        return value.isoformat()
    if value is None:
        return ""
    return str(value)


def _coerce(value: Any, target: type) -> Any:
    if target is bool:
        if isinstance(value, bool):
            return value
        text = str(value).strip().lower()
        if text in {"true", "yes", "ja", "on", "1"}:
            return True
        if text in {"false", "no", "nein", "off", "0", ""}:
            return False
        raise ValueError(value)
    if target is int:
        return int(value)
    if target is float:
        return float(value)
    return _stringify(value)


def _apply_core_properties(document: Any, config: Config, source_name: str) -> None:
    props = document.core_properties
    if config.title:
        props.title = config.title
    if config.subtitle:
        props.subject = config.subject or config.subtitle
    elif config.subject:
        props.subject = config.subject
    if config.author:
        props.author = config.author
        props.last_modified_by = config.author
    if config.keywords:
        props.keywords = config.keywords
    if config.comments:
        props.comments = config.comments
    props.category = "Markdown"
    props.modified = _dt.datetime.now()
    if source_name and source_name != "-":
        existing = props.comments or ""
        note = i18n.translate(
            config.lang, "generated_note", source=os.path.basename(source_name)
        )
        props.comments = f"{existing}\n{note}".strip()


def _starts_with_heading(html: str) -> bool:
    return html.lstrip().lower().startswith("<h1")


def _build_title_page(document: Any, config: Config) -> None:
    lookup = st.StyleLookup(document)
    for _ in range(3):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    title = document.add_paragraph(
        config.title or i18n.translate(config.lang, "untitled")
    )
    title.style = lookup["Title"]

    if config.subtitle:
        subtitle = document.add_paragraph(config.subtitle)
        subtitle.style = lookup["Subtitle"]

    for value in (config.author, config.date or _dt.date.today().isoformat()):
        if not value:
            continue
        line = document.add_paragraph(value)
        line.style = lookup[st.S_META]

    breaker = document.add_paragraph()
    breaker.paragraph_format.space_after = Pt(0)
    breaker.add_run().add_break(WD_BREAK.PAGE)


def _build_inline_title(document: Any, config: Config) -> None:
    lookup = st.StyleLookup(document)
    title = document.add_paragraph(config.title)
    title.style = lookup["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_parts = [part for part in (config.author, config.date) if part]
    if config.subtitle:
        subtitle = document.add_paragraph(config.subtitle)
        subtitle.style = lookup["Subtitle"]
        subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if meta_parts:
        line = document.add_paragraph(" · ".join(meta_parts))
        line.style = lookup[st.S_META]
        line.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _build_toc(document: Any, config: Config) -> None:
    lookup = st.StyleLookup(document)
    heading = document.add_paragraph()
    try:
        heading.style = lookup["TOCHeading"]
    except KeyError:
        heading.style = lookup["Heading1"]
    heading.add_run(config.toc_title or i18n.translate(config.lang, "toc_title"))

    body = document.add_paragraph()
    depth = max(1, min(9, config.toc_depth))
    dx.add_field(
        body,
        f'TOC \\o "1-{depth}" \\h \\z \\u',
        placeholder=i18n.translate(config.lang, "toc_placeholder"),
    )

    breaker = document.add_paragraph()
    breaker.paragraph_format.space_after = Pt(0)
    breaker.add_run().add_break(WD_BREAK.PAGE)


def _build_header_footer(document: Any, config: Config) -> None:
    for section in document.sections:
        if config.header_text:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = ""
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(config.header_text)
            run.font.size = Pt(max(7.5, config.font_size - 2))
            run.font.color.rgb = st.hex_to_rgb(config.quote_color)
            dx.set_borders(
                paragraph._p.get_or_add_pPr(),
                "w:pBdr",
                edges=("bottom",),
                size=4,
                space=2,
                color="BFBFBF",
            )

        if not (config.page_numbers or config.footer_text):
            continue

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if config.footer_text:
            run = paragraph.add_run(config.footer_text)
            run.font.size = Pt(max(7.5, config.font_size - 2))
            run.font.color.rgb = st.hex_to_rgb(config.quote_color)
            if config.page_numbers:
                separator = paragraph.add_run("   ·   ")
                separator.font.size = Pt(max(7.5, config.font_size - 2))
                separator.font.color.rgb = st.hex_to_rgb(config.quote_color)

        if config.page_numbers:
            dx.add_field(paragraph, "PAGE", placeholder="1")
            middle = paragraph.add_run(" / ")
            middle.font.size = Pt(max(7.5, config.font_size - 2))
            dx.add_field(paragraph, "NUMPAGES", placeholder="1")
            for run in paragraph.runs:
                run.font.size = Pt(max(7.5, config.font_size - 2))
                run.font.color.rgb = st.hex_to_rgb(config.quote_color)


def _remove_leading_empty_paragraph(document: Any) -> None:
    """Removes the empty leading paragraph that python-docx templates carry."""
    paragraphs = document.paragraphs
    if not paragraphs:
        return
    first = paragraphs[0]
    if first.text.strip() or first.runs:
        return
    if first.style.style_id not in ("Normal",):
        return
    if first._p.find(dx.qn("w:pPr")) is not None:
        pPr = first._p.find(dx.qn("w:pPr"))
        if pPr.find(dx.qn("w:numPr")) is not None:
            return
    parent = first._p.getparent()
    if parent is not None and len(document.paragraphs) > 1:
        parent.remove(first._p)
